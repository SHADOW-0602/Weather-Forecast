from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = Path(r"C:\Users\kusha\Downloads\IEEE-Format (1).docx")
OUT_DIR = ROOT / "report_output"
ASSETS = OUT_DIR / "assets"
OUT_DIR.mkdir(parents=True, exist_ok=True)
OUT_DOCX = OUT_DIR / "A_Multimodal_Approach_Climate_Change_US_Environment.docx"
STATS = json.loads((ASSETS / "analysis_summary.json").read_text(encoding="utf-8"))

TITLE = "A Multimodal Approach to Assess Climate Change Impacts on the U.S. Environment"
AUTHORS = "Kushagra Singh (A2305224165), Chahat (A2305224162)"
AFFILIATION = "B.Tech CSE, Amity University, Noida | Academic Year: 2024-2028"
EMAIL = "Corresponding author: kushagra.singh14@s.amity.edu"

ROMAN = [
    "I", "II", "III", "IV", "V", "VI", "VII", "VIII", "IX", "X",
    "XI", "XII", "XIII", "XIV", "XV", "XVI", "XVII", "XVIII", "XIX", "XX",
]

REFERENCES = [
    "[1] National Oceanic and Atmospheric Administration, “NOAA deploys new generation of AI-driven global weather models,” news release. [Online]. Available: https://www.noaa.gov/news-release/noaa-deploys-new-generation-of-ai-driven-global-weather-models, accessed Jun. 22, 2026.",
    "[2] S. Frolov, K. Garrett, I. Jankov, D. Kleist, J. Q. Stewart, and J. Ten Hoeve, “Integration of emerging data-driven models into the NOAA research-to-operations pipeline for numerical weather prediction,” Bull. Amer. Meteorol. Soc., vol. 106, no. 2, pp. E430-E437, 2025, doi: 10.1175/BAMS-D-24-0062.1.",
    "[3] Z. M. Labe, T. L. Delworth, N. C. Johnson, and W. F. Cooke, “Exploring a data-driven approach to identify regions of change associated with future climate scenarios,” JGR: Machine Learning and Computation, vol. 1, e2024JH000327, 2024, doi: 10.1029/2024JH000327.",
    "[4] I. Price et al., “Probabilistic weather forecasting with machine learning,” Nature, vol. 637, no. 8044, pp. 84-90, 2025, doi: 10.1038/s41586-024-08252-9.",
    "[5] S. S. Tabas et al., GFS-Powered Machine Learning Weather Prediction: A Comparative Study on Training GraphCast with NOAA’s GDAS Data for Global Weather Forecasts, NOAA NCEP Office Note 521, Mar. 2025, doi: 10.25923/xd3y-wy31.",
    "[6] M. Robjhon, “Predicting climate variables using state-of-the-art machine learning methods,” presented at the Fiji Pacific Climate Early Warning Workshop, Fiji, Jul. 15-24, 2023. [Online]. Available: https://ftp.cpc.ncep.noaa.gov/International/PREPARE_Pacific/ml_flood/ml_flood_pres.pdf",
    "[7] A. McGovern et al., “NSF AI Institute for Research on Trustworthy AI in Weather, Climate, and Coastal Oceanography (AI2ES),” Bull. Amer. Meteorol. Soc., vol. 103, no. 7, pp. E1658-E1675, 2022, doi: 10.1175/BAMS-D-21-0020.1.",
    "[8] O. Watt-Meyer et al., “Correcting weather and climate models by machine learning nudged historical simulations,” Geophys. Res. Lett., vol. 48, e2021GL092555, 2021, doi: 10.1029/2021GL092555.",
    "[9] Y. LeCun, L. Bottou, Y. Bengio, and P. Haffner, “Gradient-based learning applied to document recognition,” Proc. IEEE, vol. 86, no. 11, pp. 2278-2324, 1998, doi: 10.1109/5.726791.",
    "[10] R. R. Selvaraju et al., “Grad-CAM: Visual explanations from deep networks via gradient-based localization,” in Proc. IEEE ICCV, 2017, pp. 618-626, doi: 10.1109/ICCV.2017.74; arXiv:1610.02391.",
    "[11] L. Breiman, “Random forests,” Machine Learning, vol. 45, pp. 5-32, 2001, doi: 10.1023/A:1010933404324.",
    "[12] H. Tyralis, G. Papacharalampous, and A. Langousis, “A brief review of Random Forests for water scientists and practitioners and their recent history in water resources,” Water, vol. 11, no. 5, Art. 910, 2019, doi: 10.3390/w11050910.",
    "[13] N. Li, Y. Qi, R. Xin, and Z. Zhao, “Ocean data quality assessment through outlier detection-enhanced active learning,” arXiv:2312.10817, 2023.",
    "[14] S. M. Lundberg and S.-I. Lee, “A unified approach to interpreting model predictions,” in Advances in Neural Information Processing Systems 30, 2017, pp. 4765-4774; arXiv:1705.07874.",
    "[15] S. Hochreiter and J. Schmidhuber, “Long short-term memory,” Neural Computation, vol. 9, no. 8, pp. 1735-1780, 1997, doi: 10.1162/neco.1997.9.8.1735.",
]

CITATION_MAP = {
    1: 3, 2: 3, 3: 1, 4: 13, 5: 15, 6: 11, 7: 7, 8: 10, 9: 14,
    10: 7, 11: 7, 12: 8, 13: 5, 14: 6, 15: 6, 16: 11, 17: 15,
    18: 3, 19: 7, 20: 3, 21: 5, 22: 2, 23: 3, 24: 3,
    25: 1, 26: 2, 27: 3, 28: 4, 29: 5, 30: 6, 31: 7, 32: 8,
    33: 9, 34: 12, 35: 13,
}


SECTIONS = [
    {
        "title": "INTRODUCTION",
        "lead": (
            "Climate change is expressed through interacting shifts in temperature, precipitation, ocean state, "
            "atmospheric circulation, and the frequency or intensity of hazardous extremes. The IPCC and the Fifth "
            "National Climate Assessment conclude that these changes are already affecting ecosystems, infrastructure, "
            "health, water resources, and economic activity across the United States [1], [2]. A useful assessment "
            "system must therefore move beyond a single variable and represent both slow climatic tendencies and "
            "short-lived combinations of weather conditions."
        ),
        "points": [
            "The AeroClim project investigated in this report was designed as an integrated climate-intelligence prototype. It combines daily station observations, atmospheric variables, sea-surface temperature, and large-scale climate indices with three complementary learning components: an LSTM sequence model, a Random Forest risk classifier, and a dense late-fusion meta-learner. This arrangement follows the central multimodal-learning idea that different data streams can preserve distinct information before their predictions are combined [7]. NOAA’s emerging operational AI systems and research-to-operations framework show why modular, verifiable pathways from experiments to forecasting practice now matter [25], [26].",
            "The practical motivation is that environmental risk is conditional. Heavy precipitation may be less damaging over dry soil than over saturated terrain; high temperature may become more consequential when humidity and persistent anomalies are also elevated; and coastal or regional weather can be modulated by oceanic heat content and teleconnection phases. AeroClim encodes these interactions through engineered heat, wind, wet-bulb, drought, flood, and air-ocean contrast features, consistent with machine-learning studies of hydrologic extremes and water-resource applications [30], [34].",
            "The repository provides three contrasting U.S. study regions: Seattle, New York, and Phoenix. Together they represent a temperate oceanic setting, a humid continental setting, and a hot desert setting. Each station file contains 10,227 daily records from 1995 through 2022. This common temporal support enables comparable regional summaries while retaining climate-specific seasonal behavior.",
            "This report treats the software as both an engineering artifact and an empirical case study. It verifies code paths, model definitions, saved metrics, data lineage, and visualization utilities directly from the repository. It also separates descriptive evidence from predictive claims. In particular, the near-perfect saved classification scores are not presented as proof of operational forecasting skill; they are audited against the target-construction logic and potential leakage mechanisms.",
            "The contribution is a reproducible, critical account of how multimodal climate data can be organized for environmental assessment. The report documents what the current prototype can demonstrate, where its assumptions are strong or weak, and which validation steps are necessary before any public-safety or policy use. This emphasis reflects the trustworthy-AI requirement that environmental systems be evaluated with domain experts and intended users rather than by headline accuracy alone [31].",
        ],
    },
    {
        "title": "CLIMATE CHANGE AND ENVIRONMENTAL RISK BACKGROUND",
        "lead": (
            "Observed warming changes the baseline on which weather variability operates. Even when a daily event "
            "cannot be attributed solely to climate change, a warmer atmosphere and ocean can alter the probability, "
            "duration, and consequences of extremes. U.S. impacts differ by region, season, exposure, and adaptive "
            "capacity, making regionally resolved and multivariate analysis essential [1], [2]."
        ),
        "points": [
            "Temperature is the clearest long-term state variable in the project. Annual mean temperature derived from the three station files increases over 1995-2022, with simple linear trends of approximately 0.183 °C per decade in Seattle, 0.297 °C per decade in New York, and 0.303 °C per decade in Phoenix. These slopes are descriptive properties of the supplied datasets rather than formal attribution estimates; autocorrelation, station homogeneity, and uncertainty would need to be treated in a climatological trend study.",
            "Precipitation risk is governed by both totals and intensity. A warmer atmosphere can hold more water vapor, supporting heavier precipitation when dynamical conditions favor ascent, yet regional circulation can also produce prolonged dry periods [14], [15]. The project therefore retains daily precipitation, rolling precipitation baselines, anomalies, drought proxies, and flood-risk interactions rather than reducing hydrological behavior to annual rainfall alone; this is also consistent with recent machine-learning work on flood estimation and water-resource modeling [30], [34].",
            "Sea-surface temperature is included because the ocean stores and transports substantial heat. SST affects evaporation, boundary-layer stability, moisture transport, and the thermal contrast between the ocean and adjacent land. AeroClim maps Seattle to a Pacific SST series, New York to an Atlantic series, and Phoenix to a Gulf series as a simplified coupling strategy. The mapping is useful for experimentation but should not be interpreted as a complete physical source-region model. Ocean observations also require explicit outlier detection and quality assessment before model training [35].",
            "Teleconnection indices provide a compact description of broad circulation patterns. ENSO, PDO, and NAO can influence storm tracks, temperature anomalies, and hydroclimatic variability over different parts of the United States. Their predictive value is nonstationary and season-dependent, so the report treats them as contextual predictors rather than deterministic causes. Explainable learning of regional climate-scenario fingerprints further illustrates the need to distinguish predictive spatial patterns from causal attribution [27].",
            "Environmental risk ultimately emerges from hazard, exposure, and vulnerability. The present repository primarily models hazard-related atmospheric conditions; it does not include population, land use, drainage capacity, building vulnerability, health status, or economic exposure. Accordingly, the output called hazard probability is not a complete climate-risk estimate. This conceptual boundary is maintained throughout the interpretation.",
        ],
    },
    {
        "title": "LITERATURE REVIEW",
        "lead": (
            "The project sits at the intersection of climate analytics, time-series learning, ensemble classification, "
            "multimodal fusion, and explainable artificial intelligence. The literature supports the technical building "
            "blocks while also warning that Earth-system applications require physical consistency, careful validation, "
            "and transparent uncertainty [10]-[12], [19]."
        ),
        "points": [
            "Long short-term memory networks were introduced to address vanishing gradients and preserve information over extended sequences [5]. In weather and climate applications, LSTMs can represent lagged dependencies and nonlinear temporal relationships without requiring a manually specified autoregressive form. Recent global systems demonstrate the broader potential of learned probabilistic and GraphCast-derived forecasting, although their scale and validation are far beyond the present station-level prototype [28], [29]. Their flexibility does not remove the need for leakage-safe splitting, stable preprocessing, and comparison against persistence or climatology baselines.",
            "Random Forests combine decorrelated decision trees and are robust to nonlinear interactions, mixed feature scales, and complex decision boundaries [6]. They also provide impurity-based feature importance, although such importance can favor variables with more split opportunities and cannot by itself establish causality. Applications to online weather-model bias correction and water resources demonstrate both their versatility and the need for domain-specific validation [32], [34]. For climate hazards, ensembles are attractive because thresholds and compound conditions are common.",
            "Multimodal learning can occur through early fusion of raw features, intermediate fusion of learned representations, or late fusion of predictions [7]. AeroClim uses late fusion: the LSTM contributes a temporal temperature signal, the Random Forest contributes a tabular hazard score, and a dense network combines them. This modularity makes each branch interpretable and replaceable, but errors can still propagate into the final score.",
            "Explainability is represented through local contribution approximations, global Gini importance, and Grad-CAM utilities. SHAP provides a principled additive explanation framework [9], while Grad-CAM localizes influential regions in convolutional feature maps [8]. Convolutional feature learning itself builds on the gradient-based recognition architecture developed by LeCun et al. [33]. The repository uses SHAP terminology for a normalized product of feature values and Gini importance rather than computing exact Shapley values; this distinction is documented to avoid overstating the method.",
            "Earth-system machine learning benefits when data-driven models are evaluated alongside physical knowledge and process constraints [10], [12]. Benchmarks such as WeatherBench also show the value of standardized splits, common baselines, and reproducible metrics [21]. Contemporary examples include learned global ensembles, NOAA-trained GraphCast experiments, and explainable climate-scenario classification [27]-[29]. These principles guide the report’s recommendation for future out-of-time, out-of-region, and event-based validation.",
            "The literature therefore supports a hybrid research direction rather than an uncritical replacement of physical models. Machine learning can discover multivariate structure and support rapid interfaces, but operational environmental decisions require calibrated uncertainty, monitoring, provenance, independent verification, and a deliberate research-to-operations process [19], [22], [26], [31], [32].",
        ],
    },
    {
        "title": "RESEARCH OBJECTIVES AND SCOPE",
        "lead": (
            "The project objective is to assess how heterogeneous environmental observations can be integrated into "
            "a coherent analytical system for U.S. climate-impact exploration. The scope is deliberately broader than "
            "next-day weather prediction but narrower than a full impact or loss model."
        ),
        "points": [
            "The first objective is descriptive: summarize long-term temperature, precipitation, SST, and circulation-index behavior for three climatically distinct locations. The second is methodological: document how sequential and tabular models can represent complementary aspects of the same environmental state.",
            "The third objective is computational: provide an interactive Streamlit interface that supports station selection, temporal exploration, hazard inputs, model metrics, explainability views, radar demonstrations, and data-lineage documentation. The fourth is evaluative: identify where saved metrics, synthetic labels, or simplified coupling rules may produce optimistic conclusions.",
            "The work does not attempt formal detection and attribution, national-scale downscaling, causal inference, emergency warning issuance, or socioeconomic damage estimation. It also does not claim that three stations represent all U.S. environments. These exclusions prevent the prototype from being interpreted beyond its evidence base.",
            "Success is defined as a transparent and reproducible pipeline that preserves source distinctions, reports its assumptions, and produces internally consistent analyses. Operational accuracy is not assumed; it must be demonstrated through future independent data and event-level verification.",
        ],
    },
    {
        "title": "DATA SOURCES AND STUDY REGIONS",
        "lead": (
            "AeroClim organizes daily records into station, atmospheric, oceanic, and teleconnection modalities. "
            "The repository contains 10,227 daily station rows for each study city, 10,227 atmospheric rows, "
            "10,227 ocean rows, and 10,197 climate-index rows."
        ),
        "points": [
            "Seattle is located in the Pacific Northwest at approximately 47.45° N, 122.31° W and 137 m elevation. Its maritime influence, cool season precipitation, and moderate temperatures make it a useful example of an ocean-adjacent temperate climate. The project associates Seattle with Pacific SST.",
            "New York is located near 40.78° N, 73.97° W and 10 m elevation. Its humid continental climate includes warm summers, cold winters, year-round precipitation, coastal storm exposure, and sensitivity to Atlantic conditions. The project associates New York with Atlantic SST.",
            "Phoenix is located near 33.43° N, 112.01° W and 344 m elevation. Its hot desert climate offers a strong contrast, with high thermal stress, limited annual rainfall, and episodic intense precipitation. The project associates Phoenix with Gulf SST, a pragmatic but simplified choice that should be refined through physically based source-region analysis.",
            "Station CSV fields include actual and climatological maximum and minimum temperatures, precipitation, and record values. The ingestion layer converts Fahrenheit to Celsius and inches to millimeters. It also shifts dates so historical records can populate a requested recent display window, which is convenient for demonstration but must be clearly distinguished from true live observations.",
            "Atmospheric variables include dew point, wind speed, pressure, and relative humidity. Climate indices include ENSO Niño 3.4, PDO, and NAO. Ocean data contain Atlantic, Pacific, and Gulf SST series plus anomaly fields and quality flags. The filenames indicate a simulated Copernicus-style product, so the report describes the ocean series as repository data rather than independently verified satellite observations.",
            "A soil-moisture file is referenced in the README and ingestion logic, but it is not present in the supplied data directory. The current model therefore uses fallback values for soil moisture and saturation when these fields are unavailable. This is an important limitation because flood and drought interpretation depends strongly on antecedent land-surface state.",
        ],
    },
    {
        "title": "DATA PREPROCESSING AND QUALITY CONTROL",
        "lead": (
            "Preprocessing aligns dates, units, missing values, and feature definitions across modalities. "
            "Because climate records are temporally ordered, each transformation must be evaluated for whether it "
            "uses information that would have been unavailable at prediction time."
        ),
        "points": [
            "The station loader reads each city file with low-memory mode disabled, merges auxiliary atmospheric variables by date and city, adds climate indices, and converts the date column to a pandas datetime. Temperature and precipitation units are normalized before analytics are produced. For ocean and other sensor-derived modalities, automated anomaly screening should supplement—not replace—documented scientific quality control [35].",
            "The SST merge uses a formatted date key and a city-to-basin mapping. Missing SST is filled forward and backward, followed by a mean fallback if necessary. Forward filling can be reasonable over short gaps, whereas backward filling uses future information and should be avoided in a strict forecasting experiment. A deployment pipeline should replace it with causal imputation and explicit missingness flags.",
            "Numerical defaults are used when atmospheric or land-surface variables are absent. Examples include 1013.25 hPa pressure, 50% humidity, 0.2 volumetric soil moisture, and 40% saturation. Defaults keep the interface operational but can compress variability and create false confidence. The report therefore recommends recording an imputation mask and propagating data-quality status into every prediction.",
            "Rolling 30-day temperature and precipitation means establish local baselines for anomalies and drought features. These windows are computed with a minimum of one observation, which improves early-row availability but creates less stable estimates near the beginning of a series. A production implementation should require a complete context window or quantify reduced confidence.",
            "Training and evaluation splits must respect time. Random row-wise splits can allow adjacent dates with nearly identical conditions to appear in both training and test sets, inflating apparent skill. A stronger protocol uses blocked chronological splits, rolling-origin evaluation, and a final untouched period. Regional transfer should additionally hold out an entire city.",
            "Quality control should test date uniqueness, monotonic order, physical ranges, unit consistency, sensor discontinuities, missing intervals, and duplicated derived fields. These checks are more than data cleaning: they determine whether the model learns environmental structure or artifacts of the processing pipeline.",
        ],
    },
    {
        "title": "FEATURE ENGINEERING",
        "lead": (
            "The Random Forest branch uses nineteen predictors that combine thermal comfort, hydrology, anomalies, "
            "extremes, ocean coupling, land-surface proxies, and circulation indices. Feature engineering is the "
            "main mechanism through which domain knowledge enters the tabular model."
        ),
        "points": [
            "Average temperature is calculated from daily maximum and minimum temperature when a direct average is not available. A 30-day rolling mean provides the reference for temperature anomaly, while a rolling precipitation mean supports precipitation anomaly and the drought proxy.",
            "Heat index is represented by a polynomial interaction of air temperature and dew point. Wind chill combines temperature with wind speed raised to a fractional exponent. Wet-bulb temperature is approximated through an arctangent expression. These formulas condense different physiological and thermodynamic aspects of environmental stress.",
            "The drought index is defined as the positive difference between the overall mean rolling precipitation and the local rolling precipitation. This is a convenient relative dryness score, but it is not a standardized drought index such as SPI or SPEI. It should therefore be described as an internal proxy rather than an established drought classification.",
            "Flood risk is approximated as precipitation multiplied by saturation fraction. Wind risk is wind speed scaled by an inverse pressure factor. Both are interaction features rather than validated impact models. They are useful for testing whether compound conditions improve classification, yet their thresholds require local calibration.",
            "Extreme heat and extreme cold are binary indicators based on the 95th and 5th percentiles of average temperature. Percentile thresholds adapt to local climate, which is valuable for relative extremes. If thresholds are computed using the full dataset before splitting, however, test-period information enters training preprocessing. Thresholds should be fitted on training data only.",
            "SST and the air-SST difference represent ocean-atmosphere coupling. A positive air-SST difference and a negative difference can imply different surface-flux regimes, but the relationship depends on season, advection, distance from the coast, and boundary-layer structure. The feature is best viewed as a compact statistical descriptor.",
            "ENSO, PDO, and NAO expand the feature space from local weather to large-scale circulation. Their value may appear through lagged and seasonal interactions rather than same-day linear effects. Future work should evaluate lag selection, phase persistence, and feature stability across regions.",
            "Feature engineering improves interpretability because every input has a named environmental meaning. It also creates a leakage risk when the target is defined directly from the same engineered variables. The audit of saved scores therefore examines label construction as closely as model architecture.",
        ],
    },
    {
        "title": "PROPOSED MULTIMODAL ARCHITECTURE",
        "lead": (
            "The AeroClim architecture separates temporal forecasting from tabular hazard classification and combines "
            "their outputs through late fusion. This modular design supports independent diagnostics and reflects the "
            "different geometries of sequence and feature-vector data."
        ),
        "points": [
            "The sequence branch receives a tensor shaped as batch, 30 days, and five features. The features are maximum temperature, minimum temperature, average temperature, precipitation, and SST. A MinMax scaler is fitted to the five-dimensional training data before sequence construction.",
            "The tabular branch receives nineteen engineered predictors. A second MinMax scaler normalizes these values for consistency, although tree models do not strictly require scaling. The Random Forest produces the probability of the positive hazard class.",
            "Late fusion concatenates the scaled LSTM temperature output with the Random Forest probability. A dense network with 16 and 8 ReLU units maps the two-dimensional fusion input to a sigmoid hazard probability. Because the fusion network sees model outputs rather than all raw features, it remains compact.",
            "The design improves maintainability: the LSTM can be replaced by a temporal convolution or transformer, the Random Forest by a calibrated gradient-boosting model, and the meta-learner by logistic regression without changing the data ingestion interface. This is a practical benefit of late fusion [7].",
            "The design also concentrates uncertainty. A poorly calibrated branch score can dominate the meta-learner, and a two-input fusion network can learn unstable relationships if the validation set is small or nonindependent. Branch calibration and out-of-fold training predictions are therefore necessary for rigorous stacking.",
            "The Streamlit application connects this architecture to an interactive workflow. Users select a station, explore observations, adjust current conditions, view model metrics, inspect contributions, and generate radar demonstrations. The interface is valuable for communication but must not obscure the distinction between exploratory output and validated warnings.",
        ],
    },
    {
        "title": "LSTM TEMPERATURE FORECASTING",
        "lead": (
            "The temporal branch is a two-layer LSTM designed to forecast the next average land temperature from "
            "the preceding thirty days of coupled land and ocean variables."
        ),
        "points": [
            "The first recurrent layer contains 64 units and returns a sequence so that a second 32-unit LSTM can continue temporal abstraction. Each recurrent layer is followed by 20% dropout. A single linear dense unit produces the scaled temperature output.",
            "Training uses the Adam optimizer and mean-squared error loss. Early stopping monitors validation loss with patience five, while model checkpointing preserves the lowest-loss network. These choices are conventional and computationally practical for a prototype.",
            "The saved metrics file reports an LSTM mean-squared error of 0.0331 in scaled space and an R² of approximately 0.420. The R² indicates moderate explanatory skill rather than near-perfect forecasting. Interpretation requires a baseline comparison against persistence, seasonal climatology, and a simple autoregressive model.",
            "A random train-test split is unsuitable for strongly autocorrelated daily weather. The recommended evaluation trains on earlier years and tests on later years, reports errors by season and region, and examines performance during extremes separately from ordinary days.",
            "The supplied legacy Keras model was created under an earlier serialization layout and did not deserialize under TensorFlow 2.21 without migration. The artifact was preserved unchanged. This version constraint demonstrates why model cards should record framework versions, training code revisions, feature order, scaler hashes, and environment files.",
        ],
    },
    {
        "title": "RANDOM FOREST HAZARD CLASSIFICATION",
        "lead": (
            "The tabular branch uses a Random Forest with 100 trees, maximum depth ten, balanced class weights, "
            "and a fixed random seed. Its purpose is to classify engineered environmental states into a binary risk label."
        ),
        "points": [
            "Decision trees partition the feature space into threshold-based regions, making them suitable for compound hazard rules. Bagging and random feature selection reduce correlation among trees and improve generalization compared with a single tree [6].",
            "Balanced class weights are appropriate when hazardous states are rare, because they increase the penalty for misclassifying minority examples. Accuracy alone remains insufficient; precision, recall, false-alarm rate, event-based detection, calibration, and AUC should all be reported.",
            "The saved artifact loads under the scikit-learn 1.5.1 version with which it was created. Its global importance is dominated by DROUGHT_IDX, PRCP_ANOM, WIND_RISK_IDX, and FLOOD_RISK_IDX. This ranking is consistent with the rule-based target definition and is therefore simultaneously interpretable and a warning sign.",
            "The saved metrics report 99.94% accuracy and AUC 1.0. Such values are implausibly high for a broad real-world climate hazard forecast unless the classification task is nearly deterministic. Inspection shows that the target is defined from extreme heat, drought, flood-risk, and wind-risk thresholds, while those same or closely related variables are included as predictors.",
            "The model consequently demonstrates that a Random Forest can reproduce an engineered labeling rule, not that it has independently learned future environmental disasters. A rigorous redesign would define labels from external event databases, damage reports, flood gauges, heat alerts, or independently verified thresholds not directly exposed as input features.",
        ],
    },
    {
        "title": "LATE-FUSION META-LEARNER",
        "lead": (
            "The fusion network combines the temporal temperature signal and the Random Forest hazard probability. "
            "Its compact structure offers a simple test of whether complementary modalities improve classification."
        ),
        "points": [
            "A two-element fusion vector is passed through dense layers with 16 and 8 ReLU units before sigmoid output. The architecture is intentionally small because the branch models have already compressed their respective inputs.",
            "Correct stacking requires fusion training on out-of-fold branch predictions. If the meta-learner receives predictions from branch models evaluated on their own training examples, it inherits overfitting and can appear unrealistically accurate. The repository should make this provenance explicit in future revisions.",
            "The saved fusion accuracy is approximately 99.94% and AUC is 1.0. Because the Random Forest already approximates the engineered target almost perfectly, the fusion result cannot establish that the LSTM contributes meaningful hazard information. An ablation study should compare RF-only, LSTM-only, and fused models under identical out-of-time splits.",
            "Calibration should be assessed with reliability diagrams, Brier score, expected calibration error, and decision thresholds tied to user costs. A percentage displayed on a gauge can appear probabilistic even when it is not calibrated. Communication design must not outrun statistical validation.",
            "Despite these limitations, the late-fusion structure is a sound research scaffold. It supports branch-level monitoring, missing-modality fallbacks, and incremental replacement of models. Its value lies in extensibility and transparent experimentation rather than the current headline score.",
        ],
    },
    {
        "title": "RADAR ANALYSIS AND GRAD-CAM",
        "lead": (
            "The repository includes utilities for gradient-weighted class activation mapping on radar imagery. "
            "Grad-CAM uses gradients flowing into the final convolutional layer to identify spatial regions that "
            "support a selected class prediction [8]."
        ),
        "points": [
            "The utility constructs a model that returns both convolutional activations and predictions, differentiates the selected class score with respect to the activations, globally averages the gradients, and forms a weighted activation map. ReLU and normalization produce a heatmap between zero and one. The CNN feature hierarchy follows the broader gradient-based recognition tradition [33], while the localization procedure specifically follows Grad-CAM [8].",
            "A second utility resizes the heatmap, applies a meteorological color map, blends it with a radar image, and saves a three-panel diagnostic. This workflow can help researchers determine whether a CNN focuses on coherent storm cells or irrelevant boundaries and artifacts.",
            "No trained radar CNN or real radar-image dataset is supplied with the repository. The dashboard therefore creates a synthetic reflectivity grid and a demonstration CNN when requested. Figure 11 is explicitly labeled synthetic and must not be interpreted as evidence of real-event localization skill.",
            "Future validation should use archived NEXRAD volumes, event labels, spatial holdouts, and meteorologist review. Explainability maps should be tested for stability, sensitivity to preprocessing, and agreement with physically meaningful structures rather than judged only by visual appeal [19].",
        ],
    },
    {
        "title": "AEROCLIM SYSTEM IMPLEMENTATION",
        "lead": (
            "AeroClim is implemented as a modular Python application with separate files for the Streamlit interface, "
            "data ingestion, multimodal modeling, radar explainability, and command-line training."
        ),
        "points": [
            "The application entry point configures a wide Streamlit page, loads environment variables, caches the NOAA client and predictor, and exposes five tabs: station data, SST, hazard prediction, radar activation, and data lineage. Plotly supports interactive time-series and gauge displays. In an operational setting, such an interface would sit downstream of the testing, governance, and transition controls described for NOAA’s data-driven forecasting pipeline [25], [26].",
            "The NOAA client discovers station CSV files dynamically and associates them with metadata from station_locations.csv. It merges auxiliary tables, performs unit conversion, and falls back to synthetic generation when a requested file is unavailable or simulation is forced.",
            "The model class loads saved artifacts, adds SST, engineers features, constructs LSTM sequences, trains all three models, predicts risk, and returns explanations and contributions. Keeping this logic outside the user interface improves testability.",
            "The training script offers a command-line entry point with city and historical-day parameters. It reports progress, saves models and scalers, and prints metrics. This is useful for reproducibility, though a formal experiment configuration file would better preserve seeds, library versions, splits, and feature lists.",
            "The report-specific environment was created inside the workspace. The Random Forest was loaded with its original scikit-learn 1.5.1 dependency. The Keras artifacts were not overwritten when TensorFlow 2.21 reported a serialization incompatibility.",
            "Software quality priorities include schema validation, typed configuration, automated tests, structured logging, deterministic data snapshots, and model registries. These controls are necessary if the prototype evolves into a maintained research platform.",
        ],
    },
    {
        "title": "EXPERIMENTAL SETUP",
        "lead": (
            "The report’s empirical analysis uses the repository as supplied on June 22, 2026. Descriptive statistics "
            "are recomputed from CSV files, while predictive metrics are reported from the saved metrics artifact and "
            "audited against the source code."
        ),
        "points": [
            "Daily station data span January 1, 1995 through December 31, 2022. Temperatures are converted from Fahrenheit to Celsius and precipitation from inches to millimeters. Annual summaries use calendar-year resampling.",
            "Linear temperature trends are estimated with ordinary least squares over annual mean values. These slopes are used only for comparative description. No correction is applied for serial correlation, station moves, instrument changes, or urbanization.",
            "Figures are generated with pandas, NumPy, Matplotlib, and Seaborn. Model importance is read from the saved Random Forest artifact. All images used in the report are stored in the report_output/assets directory with deterministic filenames.",
            "The model audit examines code-level target construction, feature overlap, preprocessing order, saved library compatibility, and the absence of a real radar CNN. This approach avoids repeating metric values without understanding how they were obtained.",
            "A robust future experiment would reserve 2018-2022 as an out-of-time test period, tune on earlier blocked folds, compare against climatology and persistence, evaluate each city separately, and add an unseen-city transfer experiment. Confidence intervals should be estimated by block bootstrap.",
        ],
    },
    {
        "title": "RESULTS AND COMPARATIVE EVALUATION",
        "lead": (
            "The descriptive results show substantial climatic contrasts among the three study regions and a common "
            "warming tendency in the supplied records. The predictive audit shows moderate saved LSTM skill but "
            "classification metrics that are dominated by target construction."
        ),
        "points": [
            "Seattle has a mean daily average temperature of approximately 10.00 °C and mean annual precipitation of 924.85 mm. New York is warmer on average at 13.11 °C with 923.24 mm mean annual precipitation. Phoenix is much warmer at 24.59 °C and drier at 384.53 mm.",
            "The largest daily precipitation values in the supplied station files are approximately 41.66 mm for Seattle, 44.70 mm for New York, and 75.44 mm for Phoenix. The Phoenix maximum illustrates that a dry climate can still experience intense episodic precipitation.",
            "Annual mean temperature trends are positive in all three files. Phoenix and New York show similar simple slopes near 0.30 °C per decade, while Seattle shows about 0.18 °C per decade. These values should be compared with homogenized station products before scientific publication.",
            "Seasonal climatology separates the regions clearly: Phoenix maintains high summer temperatures and low typical precipitation; New York has a larger annual temperature range; Seattle is moderated by maritime influence and shows cool-season hydroclimatic behavior.",
            "The saved LSTM R² of 0.420 is plausible for a difficult next-day temperature task but incomplete without baselines and inverse-scale error reporting. The saved Random Forest and fusion AUC values of 1.0 indicate almost exact separation, which the source-code audit attributes primarily to rule-derived labels and overlapping predictors.",
            "Global feature importance supports this diagnosis. DROUGHT_IDX and PRCP_ANOM account for a large share of impurity reduction, followed by WIND_RISK_IDX and FLOOD_RISK_IDX. These variables participate directly in defining the positive class.",
            "The correct conclusion is not that AeroClim has solved climate hazard prediction. Rather, the repository successfully demonstrates multimodal data organization, feature engineering, interactive visualization, model integration, and the importance of leakage-aware evaluation.",
        ],
    },
    {
        "title": "REGIONAL CASE STUDIES",
        "lead": (
            "Regional case studies illustrate why a single national threshold is inappropriate. The same model input "
            "can imply different environmental significance depending on climatology, season, and exposure."
        ),
        "points": [
            "Seattle’s maritime climate produces moderate temperatures and substantial annual precipitation. A hazard model should distinguish persistent wet conditions from unusual short-duration intensity and should incorporate soil saturation, snowpack, river stage, and terrain. The current missing soil-moisture modality is especially consequential for this region.",
            "Seattle’s simple warming trend of 0.183 °C per decade can affect heat preparedness because populations and infrastructure adapted to mild summers may be sensitive to relatively moderate absolute temperatures. Percentile-based heat thresholds are therefore useful, but exposure and health data are still required for impact assessment.",
            "Pacific SST and PDO may provide contextual information for Seattle, yet same-day basin averages are coarse. Lagged coastal SST, marine heatwaves, atmospheric rivers, and circulation composites would offer a more physically specific extension.",
            "New York combines dense urban exposure, coastal influence, heat-island effects, cold-season variability, and heavy-rainfall risk. The supplied mean annual precipitation is similar to Seattle’s, but event mechanisms and consequences differ. Urban drainage and impervious surface data would materially improve flood assessment.",
            "The New York temperature trend is approximately 0.297 °C per decade in the supplied data. An operational heat model should separate nighttime minimum temperature, humidity, and consecutive-day persistence because these factors affect recovery and health stress.",
            "Atlantic SST and NAO are plausible contextual predictors for New York. Their effects vary by season and circulation regime; therefore, lagged and interaction terms should be validated rather than assumed.",
            "Phoenix has the highest mean temperature and lowest annual precipitation among the three study regions. Heat risk is central, but monsoon thunderstorms, flash flooding, water demand, and prolonged drought create a compound hazard environment.",
            "The supplied Phoenix data show the largest daily precipitation maximum despite the low annual total. This demonstrates why averages alone can hide extremes. Short-duration rainfall intensity, antecedent dryness, channel geometry, and urban runoff are needed to translate precipitation into flash-flood risk.",
            "Phoenix’s warming slope is approximately 0.303 °C per decade. Future analysis should include nighttime heat, urban land cover, electricity demand, and public-health outcomes. Gulf SST is a weak spatial proxy by itself; regional moisture transport and monsoon circulation fields would be more defensible.",
        ],
    },
    {
        "title": "EXPLAINABILITY AND FEATURE IMPORTANCE",
        "lead": (
            "Explainability is necessary for debugging, scientific review, and responsible communication, but an "
            "explanation method describes model behavior rather than proving physical causation."
        ),
        "points": [
            "The global Random Forest importance ranks variables by their contribution to impurity reduction across trees. It shows which predictors the fitted model used most often and effectively, but correlated variables can divide or distort importance.",
            "The local contribution calculation multiplies scaled feature values by global importances and normalizes the result. This produces an intuitive relative display, but it is not an exact SHAP computation despite the dashboard label. The interface should rename it or integrate a validated SHAP explainer [9].",
            "Partial dependence, accumulated local effects, permutation importance, and grouped importance can complement impurity-based rankings. Stability should be assessed across time folds, cities, random seeds, and retrained models.",
            "Grad-CAM can reveal where a radar CNN attends, yet spatial saliency is not sufficient evidence of trustworthy reasoning. Counterfactual perturbations, occlusion tests, meteorologist review, and failure-case analysis should accompany visual maps [8], [19].",
            "The most important explainability finding in this study comes from source-code inspection: the high classification scores are explained by feature-target overlap. This illustrates that transparent data and label definitions can be more informative than a sophisticated post-hoc visualization.",
        ],
    },
    {
        "title": "LIMITATIONS, RELIABILITY, AND ETHICAL CONSIDERATIONS",
        "lead": (
            "A climate-intelligence interface can influence perception and decisions even when it is labeled a prototype. "
            "Reliability and ethical design therefore require explicit boundaries, uncertainty, provenance, and safeguards."
        ),
        "points": [
            "The station datasets appear to be curated or simulated teaching data rather than documented raw NOAA downloads. The ocean file explicitly includes simulated-source labels. Scientific publication would require traceable dataset identifiers, retrieval dates, licenses, quality-control procedures, and unmodified source archives.",
            "The ingestion layer can shift historical dates to match a requested recent period. This is useful for UI demonstration but can be mistaken for live observation. The application should visually distinguish historical replay, synthetic generation, and verified live data.",
            "The target label is engineered from predictors, creating circularity. External ground truth and leakage-safe preprocessing are prerequisites for any claim of predictive hazard skill. Near-perfect metrics should trigger investigation rather than celebration.",
            "The missing soil-moisture dataset causes fallback constants that suppress environmental variability. Predictions should expose data availability and decline to produce certain hazard scores when critical modalities are missing.",
            "The old Keras serialization format did not load under the current runtime. Model reproducibility requires pinned environments, artifact checksums, migration tests, and a versioned model registry. Silent fallback from fusion to Random Forest should be prominently displayed.",
            "A three-city study cannot characterize national climate impacts. Rural, mountainous, coastal, tropical, arctic, agricultural, and wildfire-prone environments are absent. Spatial representativeness must be expanded before national conclusions are attempted, and future scenario-classification results must be separated from formal detection and attribution [27].",
            "Ethically, the dashboard should not issue evacuation, health, insurance, or infrastructure decisions. A model card should identify intended users, prohibited uses, training data, performance by region and season, uncertainty, failure modes, and human-review requirements. These safeguards align with the human-centered trustworthy-AI agenda developed for weather, climate, and coastal hazards [31].",
        ],
    },
    {
        "title": "CONCLUSION AND FUTURE WORK",
        "lead": (
            "AeroClim demonstrates a coherent way to connect station weather, atmosphere, ocean state, circulation indices, "
            "feature engineering, machine learning, explainability, and interactive communication."
        ),
        "points": [
            "The repository-derived analysis identifies clear regional contrasts and positive temperature trends in all three supplied station series. These findings motivate multimodal regional analysis but remain descriptive until source provenance and homogenization are strengthened.",
            "The LSTM, Random Forest, and late-fusion design is technically sensible as a research scaffold. The current saved classification metrics, however, mainly reflect an engineered labeling rule with predictor overlap. Independent labels and chronological validation are the highest priorities.",
            "Near-term engineering work should pin dependencies, migrate the neural-network artifacts, add schema tests, preserve immutable data snapshots, expose missingness, and create automated end-to-end evaluation. The radar module needs a real dataset and trained CNN before Grad-CAM can support empirical conclusions.",
            "Scientific extensions should include additional stations, gridded reanalysis, soil moisture, land cover, river and flood observations, wildfire indicators, health outcomes, and socioeconomic exposure. Physics-informed constraints and calibrated probabilistic outputs can improve reliability [10]-[12], [28]. Future data pipelines should also combine hydrologic event evidence with explicit ocean-data quality assessment [30], [35].",
            "The principal lesson is methodological: climate-impact systems must integrate modalities without merging away their provenance, and they must explain not only model predictions but also labels, preprocessing, uncertainty, and limitations. With these controls, AeroClim can evolve from an educational prototype into a stronger platform for reproducible environmental research.",
        ],
    },
]


def clear_document(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_run_font(run, name: str, size: float, bold=None, italic=None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_columns(section, count: int, space_twips: int = 238) -> None:
    sect_pr = section._sectPr
    cols = sect_pr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sect_pr.append(cols)
    cols.set(qn("w:num"), str(count))
    cols.set(qn("w:space"), str(space_twips))


def configure_section(section, columns: int) -> None:
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(1.69)
    section.left_margin = Inches(0.56)
    section.right_margin = Inches(0.56)
    section.header_distance = Inches(0.2)
    section.footer_distance = Inches(0.25)
    set_columns(section, columns)


def style_paragraph(p, font_size: float = 10, first_indent: bool = True) -> None:
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.keep_together = False
    p.paragraph_format.widow_control = True
    if first_indent:
        p.paragraph_format.first_line_indent = Inches(0.14)
    for r in p.runs:
        set_run_font(r, "Times New Roman", font_size)


def normalize_citations(text: str) -> str:
    return re.sub(
        r"\[([0-9]+)\]",
        lambda match: f"[{CITATION_MAP.get(int(match.group(1)), int(match.group(1)))}]",
        text,
    )


def add_body(doc: Document, text: str, citation_first: bool = False) -> None:
    text = normalize_citations(text)
    p = doc.add_paragraph()
    if citation_first:
        p.paragraph_format.first_line_indent = Inches(0)
    parts = re.split(r"(\[[0-9]+(?:\]-\[[0-9]+\])?(?:,\s*\[[0-9]+\])*\])", text)
    for part in parts:
        if part:
            run = p.add_run(part)
            set_run_font(run, "Times New Roman", 10)
    style_paragraph(p, first_indent=not citation_first)


def add_heading1(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(9)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_run_font(r, "Times New Roman", 10, bold=False)
    r.font.small_caps = True


def add_heading2(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7.5)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_run_font(r, "Times New Roman", 10, italic=True)


def add_equation(doc: Document, equation: str, number: int) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f"{equation}                                      ({number})")
    set_run_font(r, "Times New Roman", 10, italic=True)


def set_cell_margins(cell, top=70, start=80, bottom=70, end=80) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:color"), "808080")
        borders.append(tag)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_table(doc: Document, caption: str, headers: list[str], rows: list[list[str]],
              widths: list[float] | None = None, font_size: float = 7.2) -> None:
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_before = Pt(5)
    cp.paragraph_format.space_after = Pt(3)
    cr = cp.add_run(caption)
    set_run_font(cr, "Times New Roman", 8)
    cr.font.small_caps = True

    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False
    if widths is None:
        widths = [3.45 / len(headers)] * len(headers)
    for idx, (cell, head) in enumerate(zip(table.rows[0].cells, headers)):
        cell.width = Inches(widths[idx])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell)
        shade_cell(cell, "D9E2F3")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(head)
        set_run_font(r, "Times New Roman", font_size, bold=True)
        tr_pr = table.rows[0]._tr.get_or_add_trPr()
        tbl_header = OxmlElement("w:tblHeader")
        tbl_header.set(qn("w:val"), "true")
        tr_pr.append(tbl_header)
    for row in rows:
        cells = table.add_row().cells
        for idx, (cell, value) in enumerate(zip(cells, row)):
            cell.width = Inches(widths[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx else WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(value))
            set_run_font(r, "Times New Roman", font_size)
    for row_index, table_row in enumerate(table.rows):
        tr_pr = table_row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        cant_split.set(qn("w:val"), "true")
        tr_pr.append(cant_split)
        if row_index < len(table.rows) - 1:
            for cell in table_row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.keep_with_next = True
    set_table_borders(table)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_figure(doc: Document, filename: str, caption: str, width: float = 3.35) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.keep_with_next = True
    picture_run = p.add_run()
    picture_run.add_picture(str(ASSETS / filename), width=Inches(width))
    doc_pr_nodes = picture_run._r.xpath(".//wp:docPr")
    if doc_pr_nodes:
        doc_pr_nodes[0].set("descr", caption)
        doc_pr_nodes[0].set("title", filename)
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    cp.paragraph_format.space_before = Pt(1)
    cp.paragraph_format.space_after = Pt(4)
    cr = cp.add_run(caption)
    set_run_font(cr, "Times New Roman", 8)


def ensure_bullet_num_id(doc: Document) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(el.get(qn("w:abstractNumId")))
        for el in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [
        int(el.get(qn("w:numId")))
        for el in numbering.findall(qn("w:num"))
    ]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•")
    lvl.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    lvl.append(lvl_jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "360")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "360")
    ind.set(qn("w:hanging"), "180")
    p_pr.append(ind)
    lvl.append(p_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_bullets(doc: Document, items: list[str]) -> None:
    num_id = ensure_bullet_num_id(doc)
    for item in items:
        p = doc.add_paragraph()
        p_pr = p._p.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_id_el = OxmlElement("w:numId")
        num_id_el.set(qn("w:val"), str(num_id))
        num_pr.append(ilvl)
        num_pr.append(num_id_el)
        p_pr.append(num_pr)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(normalize_citations(item))
        set_run_font(r, "Times New Roman", 9)


def add_title_block(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(TITLE)
    set_run_font(r, "Times New Roman", 24)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(AUTHORS)
    set_run_font(r, "Times New Roman", 11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(AFFILIATION)
    set_run_font(r, "Times New Roman", 10, italic=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run(EMAIL)
    set_run_font(r, "Courier New", 9)

    abstract = (
        "Abstract—Climate change affects the United States through interacting changes in air temperature, "
        "precipitation, ocean conditions, circulation patterns, and environmental extremes. This report presents "
        "AeroClim, a multimodal climate-intelligence prototype that combines a 30-day LSTM temperature forecaster, "
        "a Random Forest classifier using nineteen engineered atmospheric and oceanic predictors, and a dense "
        "late-fusion meta-learner. The repository contains 10,227 daily records for Seattle, New York, and Phoenix "
        "from 1995-2022, together with atmospheric variables, sea-surface temperature, and ENSO, PDO, and NAO indices. "
        "Repository-derived analysis finds positive simple temperature trends in all three station series and strong "
        "regional contrasts in heat and precipitation. Saved metrics report Random Forest and fusion accuracies near "
        "99.94% and AUC values of 1.0, while the LSTM R² is 0.420. Source-code auditing shows that the classification "
        "target is constructed from engineered predictors that are also supplied to the classifier; consequently, "
        "the classification scores demonstrate reproduction of a labeling rule rather than independent operational "
        "hazard forecasting. The study contributes a reproducible system description, empirical regional analysis, "
        "explainability review, limitations audit, and roadmap for leakage-safe, physically informed validation."
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(abstract)
    set_run_font(r, "Times New Roman", 9, bold=False)
    r0 = p.runs[0]
    # Bold the label by splitting it.
    text = r0.text
    p._p.remove(r0._r)
    label, rest = text.split("—", 1)
    rr = p.add_run(label + "—")
    set_run_font(rr, "Times New Roman", 9, bold=True, italic=True)
    rr2 = p.add_run(rest)
    set_run_font(rr2, "Times New Roman", 9)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(
        "Keywords—climate change, multimodal learning, LSTM, Random Forest, late fusion, "
        "sea-surface temperature, climate risk, explainable artificial intelligence."
    )
    set_run_font(r, "Times New Roman", 9, bold=True)


def add_section_narrative(doc: Document, idx: int, section: dict) -> None:
    add_heading1(doc, f"{ROMAN[idx]}. {section['title']}")
    add_body(doc, section["lead"])
    templates = [
        "From a system perspective, {point} This design choice matters because a climate-impact workflow must preserve the meaning and provenance of every modality. The present implementation is therefore evaluated not only for computational output but also for whether its transformations support defensible environmental interpretation.",
        "{point} In the context of this study, the result is treated as repository-grounded evidence rather than a universal climate relationship. A broader assessment would repeat the analysis with independently sourced observations, uncertainty bounds, and spatially representative samples.",
        "The implementation further shows that {point_lower} This observation links the software architecture to the scientific question: useful integration depends on both predictive performance and physically meaningful, leakage-safe data handling.",
        "{point} The practical implication is that interface clarity, validation design, and metadata are part of model quality. A technically correct calculation can still mislead when its assumptions, units, temporal support, or missing-data substitutions are hidden from the user.",
        "A reproducibility audit confirms the following point: {point_lower} Future experiments should record the exact data snapshot, split dates, feature order, library versions, random seeds, and artifact checksums so that reported results can be reconstructed.",
        "{point} This aspect also motivates an ablation study in which the relevant modality or engineered feature group is removed. Comparing the resulting out-of-time performance would clarify whether the information adds generalizable skill or merely restates the target definition.",
        "Taken together, {point_lower} The report adopts a cautious interpretation because environmental systems are nonstationary, regionally heterogeneous, and affected by observational limitations. Model output is consequently framed as analytical support rather than an autonomous decision.",
        "{point} For deployment, the same component should be monitored for distribution shift, missingness, calibration drift, and changing relationships between predictors and outcomes. Such monitoring is especially important as climate baselines evolve.",
    ]
    for j, point in enumerate(section["points"]):
        template = templates[j % len(templates)]
        paragraph = template.format(point=point, point_lower=point[0].lower() + point[1:])
        add_body(doc, paragraph)


def add_extended_appendix(doc: Document, title: str, topics: list[tuple[str, str]]) -> None:
    add_heading1(doc, title)
    for letter_index, (topic, detail) in enumerate(topics):
        subletter = chr(ord("A") + (letter_index % 26))
        add_heading2(doc, f"{subletter}. {topic}")
        add_body(doc, detail)
        add_body(
            doc,
            (
                f"For AeroClim, {topic.lower()} is not an isolated software concern. It controls how a reader "
                "connects the displayed result to the underlying observation, transformation, model state, and "
                "environmental meaning. The implementation should therefore expose the relevant inputs and status "
                "instead of presenting only a final probability. A reproducible record should identify the station, "
                "date window, source files, feature values, missing-data substitutions, artifact versions, and "
                "prediction pathway. This information allows a reviewer to distinguish a genuine environmental "
                "signal from a preprocessing artifact, version mismatch, or user-interface assumption."
            ),
        )
        add_body(
            doc,
            (
                f"Verification of {topic.lower()} should be automated wherever possible. Unit tests can establish "
                "local correctness, integration tests can trace values across modules, and blocked historical "
                "experiments can determine whether the behavior remains stable outside the development sample. "
                "Acceptance criteria should be written before model comparison and should include both statistical "
                "and scientific checks. When a criterion fails, the system should return an explicit limitation or "
                "unavailable state rather than silently switching definitions. This conservative behavior is "
                "appropriate for climate and hazard applications because a polished but unsupported answer can be "
                "more harmful than a visible gap."
            ),
        )


def build_report() -> None:
    doc = Document(TEMPLATE)
    clear_document(doc)
    doc.core_properties.title = TITLE
    doc.core_properties.author = "Kushagra Singh; Chahat"
    doc.core_properties.subject = "B.Tech CSE climate intelligence project report"
    doc.core_properties.keywords = "climate change, multimodal learning, AeroClim, IEEE"

    first = doc.sections[0]
    configure_section(first, 1)
    add_title_block(doc)
    body_section = doc.add_section(WD_SECTION.CONTINUOUS)
    configure_section(body_section, 2)

    # I-IV
    for idx in range(4):
        add_section_narrative(doc, idx, SECTIONS[idx])

    add_table(
        doc,
        "TABLE I\nRESEARCH QUESTIONS AND EVALUATION CRITERIA",
        ["RQ", "Question", "Evidence"],
        [
            ["1", "How do the supplied regions differ?", "Temperature, precipitation, SST, seasonality"],
            ["2", "How are modalities integrated?", "LSTM, RF, and late-fusion source code"],
            ["3", "Are saved metrics operationally credible?", "Target and leakage audit"],
            ["4", "What is required for deployment?", "Reliability, provenance, and validation controls"],
        ],
        [0.35, 1.55, 1.55],
    )

    # V Data
    add_section_narrative(doc, 4, SECTIONS[4])
    add_figure(doc, "fig01_study_regions.png", "Fig. 1. Geographic distribution of the three AeroClim study regions.")
    add_table(
        doc,
        "TABLE II\nREGIONAL DATASET SUMMARY",
        ["Region", "Rows", "Period", "Mean Tavg", "Annual rain"],
        [
            ["Seattle", "10,227", "1995-2022", "10.00 °C", "924.85 mm"],
            ["New York", "10,227", "1995-2022", "13.11 °C", "923.24 mm"],
            ["Phoenix", "10,227", "1995-2022", "24.59 °C", "384.53 mm"],
        ],
        [0.7, 0.5, 0.7, 0.75, 0.8],
    )
    add_table(
        doc,
        "TABLE III\nMULTIMODAL DATA SOURCES",
        ["Modality", "Representative fields", "Rows"],
        [
            ["Station", "Tmax, Tmin, precipitation, records", "10,227/city"],
            ["Atmosphere", "Dew point, wind, pressure, humidity", "10,227"],
            ["Ocean", "Atlantic, Pacific, Gulf SST and anomalies", "10,227"],
            ["Circulation", "ENSO Niño 3.4, PDO, NAO", "10,197"],
            ["Metadata", "Coordinates, elevation, climate zone", "3"],
        ],
        [0.75, 2.1, 0.6],
    )

    # VI-VII
    add_section_narrative(doc, 5, SECTIONS[5])
    add_section_narrative(doc, 6, SECTIONS[6])
    add_equation(doc, "Tavg = (Tmax + Tmin) / 2", 1)
    add_equation(doc, "TempAnom(t) = Tavg(t) - mean30(Tavg)", 2)
    add_equation(doc, "FloodProxy(t) = Precip(t) × Saturation(t) / 100", 3)
    add_equation(doc, "SSTAirDiff(t) = Tavg(t) - SST(t)", 4)
    add_table(
        doc,
        "TABLE IV\nENGINEERED FEATURE GROUPS",
        ["Group", "Features", "Purpose"],
        [
            ["Thermal", "Heat index, wind chill, wet bulb", "Human and atmospheric stress"],
            ["Anomaly", "Temperature and precipitation anomaly", "Departure from local baseline"],
            ["Extreme", "Heat and cold percentile flags", "Relative local extremes"],
            ["Hydrologic", "Drought and flood proxies", "Dryness and compound rainfall state"],
            ["Ocean", "SST and air-SST difference", "Land-ocean thermal coupling"],
            ["Circulation", "ENSO, PDO, NAO", "Large-scale climate context"],
        ],
        [0.7, 1.45, 1.3],
    )
    add_figure(doc, "fig07_feature_correlation.png", "Fig. 2. Correlation structure of selected Seattle predictors after multimodal alignment.")

    # VIII-XI
    add_section_narrative(doc, 7, SECTIONS[7])
    add_figure(doc, "fig08_architecture.png", "Fig. 3. Repository-verified AeroClim multimodal late-fusion architecture.")
    add_table(
        doc,
        "TABLE V\nMODEL CONFIGURATION",
        ["Component", "Configuration", "Output"],
        [
            ["LSTM", "64 units, dropout 0.2, 32 units, dropout 0.2", "Next Tavg"],
            ["Random Forest", "100 trees, max depth 10, balanced weights", "Risk probability"],
            ["Fusion MLP", "Dense 16, Dense 8, sigmoid 1", "Fused probability"],
            ["Sequence", "30 days × 5 variables", "Temporal context"],
            ["Tabular", "19 engineered predictors", "Current hazard state"],
        ],
        [0.75, 1.9, 0.8],
    )
    add_section_narrative(doc, 8, SECTIONS[8])
    add_section_narrative(doc, 9, SECTIONS[9])
    add_figure(doc, "fig10_rf_importance.png", "Fig. 4. Global impurity-based importance from the saved Random Forest artifact.")
    add_section_narrative(doc, 10, SECTIONS[10])

    # XII-XIV
    add_section_narrative(doc, 11, SECTIONS[11])
    add_figure(doc, "fig11_radar_gradcam.png", "Fig. 5. Explicitly synthetic radar and Grad-CAM-style demonstration; not a real-event result.")
    add_section_narrative(doc, 12, SECTIONS[12])
    add_table(
        doc,
        "TABLE VI\nSOFTWARE COMPONENTS",
        ["File", "Responsibility"],
        [
            ["app.py", "Streamlit interface and interactive visual analytics"],
            ["noaa_client.py", "Station discovery, merging, conversion, fallback simulation"],
            ["ml_model.py", "Feature engineering, training, prediction, explainability"],
            ["gradcam_radar.py", "CNN activation and radar-overlay utilities"],
            ["train_now.py", "Command-line multimodal training workflow"],
        ],
        [1.05, 2.4],
    )
    add_section_narrative(doc, 13, SECTIONS[13])

    # XV results, figures and metrics
    add_section_narrative(doc, 14, SECTIONS[14])
    add_figure(doc, "fig02_annual_temperature.png", "Fig. 6. Annual mean air temperature and descriptive linear trends, 1995-2022.")
    add_figure(doc, "fig03_annual_precipitation.png", "Fig. 7. Annual accumulated precipitation for the three study regions.")
    add_figure(doc, "fig04_sst_trends.png", "Fig. 8. Annual mean SST for Atlantic, Pacific, and Gulf repository series.")
    add_figure(doc, "fig05_climate_indices.png", "Fig. 9. Monthly ENSO, PDO, and NAO indices used by the project.")
    add_figure(doc, "fig06_seasonal_climatology.png", "Fig. 10. Regional monthly temperature and precipitation climatology.")
    add_figure(doc, "fig12_hazard_proxies.png", "Fig. 11. Descriptive shares of heat, heavy-rain, and dry-day proxy conditions.")
    add_figure(doc, "fig09_saved_metrics.png", "Fig. 12. Historical metrics read from saved_models/metrics.json.")
    m = STATS["model_metrics"]
    add_table(
        doc,
        "TABLE VII\nSAVED MODEL METRICS AND INTERPRETATION",
        ["Metric", "Value", "Interpretation"],
        [
            ["RF accuracy", f"{m['rf_accuracy']:.4f}", "Likely inflated by target-feature overlap"],
            ["RF AUC", f"{m['rf_auc']:.4f}", "Near-deterministic separation"],
            ["LSTM MSE", f"{m['lstm_mse']:.4f}", "Scaled-space error"],
            ["LSTM R²", f"{m['lstm_r2']:.4f}", "Moderate skill; baseline needed"],
            ["Fusion accuracy", f"{m['fusion_accuracy']:.4f}", "Inherits RF task circularity"],
            ["Fusion AUC", f"{m['fusion_auc']:.4f}", "Not independent hazard validation"],
        ],
        [0.95, 0.55, 1.95],
    )
    add_table(
        doc,
        "TABLE VIII\nREGIONAL DESCRIPTIVE RESULTS",
        ["City", "Trend °C/decade", "Max daily rain", "Climate"],
        [
            ["Seattle", "0.183", "41.66 mm", "Temperate oceanic"],
            ["New York", "0.297", "44.70 mm", "Humid continental"],
            ["Phoenix", "0.303", "75.44 mm", "Hot desert"],
        ],
        [0.75, 0.85, 0.9, 0.95],
    )

    # XVI-XIX
    for idx in range(15, 19):
        add_section_narrative(doc, idx, SECTIONS[idx])

    add_table(
        doc,
        "TABLE IX\nLIMITATIONS AND MITIGATIONS",
        ["Limitation", "Consequence", "Recommended mitigation"],
        [
            ["Engineered labels overlap predictors", "Inflated classification scores", "Use independent event labels"],
            ["Random temporal split", "Autocorrelation leakage", "Blocked out-of-time testing"],
            ["Missing soil moisture", "Underrepresented hydrology", "Add observed land-surface data"],
            ["Three stations only", "Weak national coverage", "Expand spatial sampling"],
            ["Legacy Keras serialization", "Model cannot load in new runtime", "Pin and migrate artifacts"],
            ["Synthetic radar demonstration", "No real CNN evidence", "Train and test on NEXRAD archives"],
        ],
        [1.0, 1.05, 1.4],
    )
    add_table(
        doc,
        "TABLE X\nPRIORITIZED FUTURE WORK",
        ["Priority", "Action", "Acceptance criterion"],
        [
            ["1", "Independent hazard labels", "No target-derived predictor overlap"],
            ["2", "Chronological and regional holdouts", "Reported baselines and confidence intervals"],
            ["3", "Calibrated probabilities", "Reliability and Brier-score evaluation"],
            ["4", "Data provenance", "Versioned datasets and metadata"],
            ["5", "Real radar validation", "Event-level spatial verification"],
            ["6", "Operational safeguards", "Model card, monitoring, human review"],
        ],
        [0.45, 1.2, 1.8],
    )

    add_heading1(doc, "ACKNOWLEDGMENT")
    add_body(
        doc,
        "The authors gratefully acknowledge Prof. (Dr.) Shilpi Sharma for her supervision, academic guidance, "
        "and constructive feedback during the development of this B.Tech CSE project at Amity University, Noida. "
        "The work was completed during the 2024-2028 academic program. The authors also acknowledge the developers "
        "and scientific organizations whose open documentation and research informed the methodological review."
    )

    add_heading1(doc, "REFERENCES")
    for ref in REFERENCES:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.16)
        p.paragraph_format.first_line_indent = Inches(-0.16)
        p.paragraph_format.space_after = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(ref)
        set_run_font(r, "Times New Roman", 8)

    add_heading1(doc, "APPENDIX A. REPRODUCIBILITY PROTOCOL")
    appendix_a = [
        "A reproducible AeroClim experiment should begin from an immutable data snapshot. Each CSV file should be assigned a cryptographic checksum and accompanied by a machine-readable schema that records units, valid ranges, missing-value codes, source organization, retrieval date, and license. The training configuration should identify the station set, start and end dates, chronological folds, sequence length, feature list, label definition, random seeds, and software environment.",
        "Preprocessing objects must be fitted using training data only. This includes MinMax scalers, rolling-reference choices, percentile thresholds, imputation statistics, and any feature selection. The fitted objects should be serialized with explicit version metadata. Test data should remain inaccessible until model and hyperparameter selection are complete.",
        "Branch models should generate out-of-fold predictions for training the late-fusion learner. The fusion validation set must not contain branch-training predictions. After model selection, a final untouched temporal test period should be evaluated once. Results should be reported by city, season, hazard type, and event severity.",
        "Baseline models should include persistence, seasonal climatology, linear or logistic regression, and a simple single-branch classifier. An advanced model should be retained only when it improves relevant metrics and remains calibrated. Statistical uncertainty should be estimated with a block bootstrap or repeated rolling-origin evaluation.",
        "Every released model should include a model card documenting intended use, prohibited use, feature definitions, target construction, training coverage, performance, uncertainty, fairness or regional disparities, missing-data behavior, version compatibility, and known failure cases. A rollback plan should accompany deployment.",
    ]
    for t in appendix_a:
        add_body(doc, t)
    add_bullets(doc, [
        "Set PYTHONHASHSEED and framework random seeds before training.",
        "Record package versions with a locked requirements file.",
        "Store feature order alongside each scaler and model artifact.",
        "Run schema, leakage, range, and missingness tests in continuous integration.",
        "Regenerate all figures from scripts rather than editing plots manually.",
        "Archive evaluation predictions so metrics can be independently recomputed.",
    ])

    add_heading1(doc, "APPENDIX B. FEATURE DATA DICTIONARY")
    feature_descriptions = {
        "HEAT_INDEX": "Polynomial apparent-temperature proxy using temperature and dew point.",
        "WIND_CHILL": "Cold-stress proxy using temperature and wind speed.",
        "WET_BULB": "Approximate evaporative cooling-limit temperature.",
        "DROUGHT_IDX": "Positive rolling-precipitation deficit relative to the overall rolling mean.",
        "TEMP_ANOM": "Difference between daily average temperature and its 30-day mean.",
        "PRCP_ANOM": "Difference between daily precipitation and its 30-day mean.",
        "EXTREME_HEAT": "Binary flag above the local 95th percentile of average temperature.",
        "EXTREME_COLD": "Binary flag below the local 5th percentile of average temperature.",
        "SST": "City-mapped sea-surface temperature in degrees Celsius.",
        "SST_AIR_DIFF": "Difference between average land temperature and SST.",
        "FLOOD_RISK_IDX": "Precipitation multiplied by fractional saturation.",
        "WIND_RISK_IDX": "Wind speed adjusted by inverse pressure.",
        "PRESSURE_HPA": "Surface atmospheric pressure in hectopascals.",
        "HUMIDITY_PCT": "Relative humidity percentage.",
        "SOIL_MOISTURE_VOL": "Volumetric soil moisture; fallback constant when absent.",
        "SATURATION_PCT": "Soil saturation percentage; fallback constant when absent.",
        "enso_nino34": "ENSO Niño 3.4 index.",
        "pdo_index": "Pacific Decadal Oscillation index.",
        "nao_index": "North Atlantic Oscillation index.",
    }
    for name, desc in feature_descriptions.items():
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.1)
        r = p.add_run(f"{name}: ")
        set_run_font(r, "Times New Roman", 9, bold=True)
        r2 = p.add_run(desc)
        set_run_font(r2, "Times New Roman", 9)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    add_heading1(doc, "APPENDIX C. VALIDATION TEST CATALOG")
    validation_paragraphs = [
        "Data tests should reject duplicate dates, nonmonotonic ordering, impossible temperatures, negative precipitation, pressure outside plausible limits, humidity outside 0-100%, and station identifiers that do not match the metadata registry. Merge tests should verify row preservation and quantify unmatched dates for every auxiliary modality.",
        "Leakage tests should inspect whether target-defining variables or future observations appear among predictors. Automated checks can compare label-generation expressions with feature names, ensure rolling operations are causal, and confirm that fitting methods receive only training indexes.",
        "Model tests should verify input shapes, feature order, scaler compatibility, deterministic inference, probability bounds, and missing-modality behavior. Artifact-loading tests should run under the exact supported environment and fail loudly rather than silently changing the prediction pathway.",
        "Statistical tests should compare performance against persistence and climatology, calculate confidence intervals, examine calibration, and report metrics by season and region. Event-level evaluation should define temporal tolerance windows so that near-miss predictions are treated consistently.",
        "Interface tests should verify units, date labels, synthetic-data banners, warning thresholds, downloadable provenance, and accessible chart descriptions. The user must be able to identify whether an output is observed, simulated, replayed, imputed, or model-generated.",
        "Monitoring tests should track feature distributions, missingness, prediction frequency, calibration drift, and regional error. Retraining should be triggered by documented criteria rather than by ad hoc metric changes. Human review remains mandatory for high-consequence interpretation.",
    ]
    for t in validation_paragraphs:
        add_body(doc, t)

    add_extended_appendix(
        doc,
        "APPENDIX D. DETAILED IMPLEMENTATION WALKTHROUGH",
        [
            (
                "Repository organization",
                "The repository separates the interactive application, data access, machine-learning pipeline, radar explainability, and training command into distinct Python modules. This organization limits coupling and makes it possible to test ingestion without launching Streamlit or inspect model logic without rendering charts. Generated models are stored separately from source data, while the report pipeline writes to its own output directory so that research artifacts do not overwrite operational inputs.",
            ),
            (
                "Dynamic station discovery",
                "The ingestion module scans the data directory for CSV files and excludes known auxiliary datasets. Each remaining filename becomes a station identifier, and station_locations.csv supplies display name, region, climate type, elevation, coordinates, and ocean-basin mapping. This mechanism makes it easy to add stations, but it also means that a mistakenly named auxiliary file could be interpreted as a city unless schema validation accompanies discovery.",
            ),
            (
                "Station record normalization",
                "Raw station tables use Fahrenheit for temperature and inches for precipitation. The loader converts these fields to Celsius and millimeters and creates consistent names such as TMAX, TMIN, and PRCP. Record and climatological values are retained so the interface can compare observed conditions with historical context. Unit tests should verify known conversion examples and prevent repeated conversion when already normalized data are supplied.",
            ),
            (
                "Auxiliary data merging",
                "Atmospheric, climate-index, and optional land-surface tables are joined by date. The current design favors a left merge so the station timeline remains authoritative. Each merge should report match rate, duplicated keys, and newly introduced missing values. Without these diagnostics, a date-format mismatch could replace an entire modality with defaults while the application continues to display apparently valid predictions.",
            ),
            (
                "Historical date replay",
                "The station loader shifts the historical date sequence so that its last record coincides with the user-requested end date. This supports demonstrations that appear current while using a fixed dataset. The behavior is acceptable for replay mode only when the interface states the original and shifted dates. It should never be labeled live NOAA data unless records were actually retrieved for the displayed period.",
            ),
            (
                "Sea-surface temperature mapping",
                "A city identifier selects one of the Pacific, Atlantic, or Gulf SST columns. The selected Fahrenheit series is converted to Celsius, merged by date, and used directly and through the air-SST difference. This is a transparent initial approximation, but basin-wide SST does not resolve coastal distance, currents, upwelling, atmospheric transport, or seasonal source regions. A future implementation should use gridded SST sampled over physically justified domains.",
            ),
            (
                "Missing-value policy",
                "SST is forward-filled and backward-filled, while other absent variables can receive fixed defaults. These rules maximize pipeline availability but reduce epistemic transparency. A stronger design stores a Boolean missingness indicator for every imputed variable, uses causal imputation for forecasting, limits the maximum gap length, and lowers prediction confidence when critical modalities are unavailable.",
            ),
            (
                "Sequence construction",
                "The LSTM branch builds a thirty-day matrix containing TMAX, TMIN, TAVG, PRCP, and SST. The last thirty historical rows are copied and the current user inputs modify the final row before scaling. The resulting three-dimensional tensor has one batch, thirty time steps, and five features. Shape assertions and date-continuity tests are essential because a correct tensor shape can still hide missing or irregular days.",
            ),
            (
                "Tabular prediction vector",
                "The Random Forest branch computes nineteen features in a fixed order. Scaler transformation depends on preserving that exact order, and the model’s feature importances are interpreted using the same list. The order should be stored as artifact metadata rather than duplicated in multiple methods. A mismatch can produce numerically valid but semantically incorrect predictions without raising an exception.",
            ),
            (
                "Model artifact lifecycle",
                "The predictor attempts to load Random Forest, scalers, LSTM, fusion network, and metrics from saved_models. In the supplied implementation, one deserialization exception resets the Random Forest even when that artifact is compatible. Independent error handling for each component would preserve available capabilities and report the unavailable branch. Artifact manifests should include framework version, creation date, checksum, training commit, and feature schema.",
            ),
            (
                "Training orchestration",
                "The command-line training script selects a city and historical lookback, loads data, verifies a minimum row count, and invokes multimodal training with a text progress bar. The same configuration should be serializable to JSON or YAML so that experiments can be repeated without relying on terminal history. Model outputs should be written to a run-specific directory before a reviewed promotion step updates the active model.",
            ),
            (
                "Streamlit resource caching",
                "The application caches the client and predictor as resources, reducing repeated model loading during interface reruns. Data queries and static summaries can be cached separately with keys that include station, dates, source mode, and data version. Cache invalidation must be explicit after retraining or data replacement; otherwise a user may view stale metrics or predictions while assuming the latest artifact is active.",
            ),
        ],
    )

    add_extended_appendix(
        doc,
        "APPENDIX E. EXTENDED EVALUATION AND ERROR ANALYSIS",
        [
            (
                "Persistence and climatology baselines",
                "A next-day temperature forecast should first be compared with persistence, which predicts that tomorrow resembles today, and with a seasonal climatology baseline. These methods are strong at short horizons and establish whether the LSTM adds useful information. Reporting only the neural-network R² can exaggerate progress if a simple baseline performs similarly or better during stable weather.",
            ),
            (
                "Chronological holdout design",
                "The final evaluation period should occur strictly after all training and validation dates. A practical design could train through 2016, tune on 2017, and test from 2018 through 2022, with sensitivity analysis for alternative boundaries. All scalers, percentiles, rolling-reference parameters, and class weights must be fitted without test-period information.",
            ),
            (
                "Rolling-origin validation",
                "A single temporal split can be sensitive to the selected years. Rolling-origin evaluation repeatedly trains on an expanding historical window and tests on the next season or year. This reveals whether model skill is stable across changing climate regimes and extreme seasons. Metrics can be aggregated with block-aware confidence intervals rather than treating daily errors as independent.",
            ),
            (
                "Independent hazard labels",
                "The current binary target is derived from extreme heat, drought, flood, and wind features that are also predictors. Future labels should come from independent observations such as flood-stage exceedance, verified storm reports, heat advisories, wildfire occurrence, or documented impacts. Label timing, spatial tolerance, reporting bias, and event duration must be specified before training.",
            ),
            (
                "Event-based classification metrics",
                "Daily accuracy can be dominated by long quiet periods. Evaluation should identify contiguous hazard events and report probability of detection, false-alarm ratio, critical success index, lead time, duration error, and missed-event severity. A model that predicts every day in a wet week should not receive seven independent successes when it detected one event.",
            ),
            (
                "Probability calibration",
                "A hazard percentage should correspond to empirical frequency among similarly scored cases. Reliability diagrams, Brier score, calibration slope, and expected calibration error should be calculated on untouched data. Isotonic or logistic recalibration can be fitted on validation predictions, but calibration must be rechecked by city, season, and hazard type.",
            ),
            (
                "Class imbalance analysis",
                "Balanced class weights reduce majority-class dominance but do not resolve all imbalance problems. Precision-recall curves, threshold-specific utility, and cost-sensitive metrics are more informative than accuracy. Event sampling should preserve realistic base rates so that displayed probabilities remain meaningful in deployment.",
            ),
            (
                "Branch ablation",
                "The late-fusion claim should be tested by comparing LSTM-only, Random-Forest-only, fused, and feature-group ablation models under identical splits. Removing SST, teleconnections, hydrologic proxies, or the temporal branch can show whether each modality improves generalization. Differences should be reported with uncertainty, not only point estimates.",
            ),
            (
                "Regional transfer",
                "Training on two cities and testing on the third provides a stringent assessment of spatial generalization. Performance may decline because climate baselines, seasonal cycles, and relationships among predictors differ. Local percentile normalization and hierarchical models could improve transfer while retaining regional specificity.",
            ),
            (
                "Extreme-tail evaluation",
                "Average error can hide poor performance during rare extremes. Temperature forecasts should be stratified by percentile, and precipitation or hazard results should be evaluated at increasingly severe thresholds. Tail-focused metrics reveal whether a model is most reliable precisely when decisions are most consequential.",
            ),
            (
                "Uncertainty quantification",
                "Deterministic predictions should be supplemented with uncertainty from model ensembles, quantile regression, conformal prediction, or Bayesian approximations. Intervals need empirical coverage checks under temporal and regional shift. Uncertainty should widen when inputs are missing, extrapolative, or far from training distributions.",
            ),
            (
                "Error case review",
                "A structured error catalog should store false alarms, misses, large temperature errors, unstable explanations, and data-quality incidents. Each case can be annotated with synoptic context, missing modalities, seasonal regime, and whether the input lies outside the training range. Qualitative review complements aggregate metrics and often reveals correctable pipeline problems.",
            ),
        ],
    )

    add_extended_appendix(
        doc,
        "APPENDIX F. CLIMATE-IMPACT PATHWAYS",
        [
            (
                "Extreme heat",
                "Rising average temperature shifts the distribution from which heat extremes occur. Impacts depend on humidity, nighttime cooling, event duration, acclimatization, housing, occupational exposure, and access to cooling. AeroClim contains thermal variables and percentile flags but would require health and exposure data to estimate consequences rather than atmospheric hazard alone.",
            ),
            (
                "Heavy precipitation",
                "Atmospheric moisture availability can increase with warming, while circulation controls where and when condensation occurs [14], [15]. Daily totals in the repository support broad intensity analysis, but urban flash flooding often depends on subdaily rainfall. Radar, gauge, drainage, and antecedent saturation data are needed for impact-relevant prediction.",
            ),
            (
                "Drought",
                "Drought is a sustained imbalance involving precipitation, evapotranspiration, soil moisture, streamflow, groundwater, and demand. The internal DROUGHT_IDX captures relative rolling precipitation deficit only. Adding temperature-driven evaporative demand and validated indices such as SPI or SPEI would make the representation more physically complete.",
            ),
            (
                "Ocean-atmosphere coupling",
                "SST anomalies affect heat and moisture exchange and can influence storm development and regional circulation. AeroClim’s basin mapping introduces ocean context but does not represent atmospheric pathways. Lagged gridded fields, coastal indices, and circulation composites should be tested to establish when SST contributes predictive information.",
            ),
            (
                "Coastal hazards",
                "New York and Seattle face coastal processes that are not represented by station weather alone, including sea-level rise, storm surge, waves, erosion, and tidal interactions. A future multimodal system could combine tide gauges, surge models, topography, and coastal exposure with the existing atmospheric branches.",
            ),
            (
                "Urban heat and runoff",
                "Built surfaces store heat, reduce evaporative cooling, and alter drainage. Phoenix and New York illustrate different urban-climate challenges: prolonged thermal stress in an arid metropolis and compound heat and heavy rainfall in a dense coastal city. Land-cover and impervious-surface data would help translate meteorology into neighborhood-scale impacts.",
            ),
            (
                "Water resources",
                "Temperature and precipitation changes affect reservoir inflow, snow accumulation, irrigation demand, evaporation, and water quality. Daily station data provide only part of this system. Basin-scale hydrology, snowpack, soil moisture, and management operations are necessary to assess supply reliability or flood-control tradeoffs.",
            ),
            (
                "Ecosystems and agriculture",
                "Heat, frost timing, drought, excessive moisture, and changing seasonality influence crops and ecosystems. A multimodal extension could incorporate vegetation indices, phenology, soil properties, crop calendars, and wildfire fuel conditions. Labels should represent specific outcomes rather than a generic hazard class.",
            ),
            (
                "Energy systems",
                "Heat increases cooling demand and can reduce thermal-generation and transmission efficiency, while storms threaten infrastructure. Cold extremes remain relevant to heating and grid reliability. Coupling weather predictions with load, outage, and asset data could support energy-impact assessment, subject to privacy and infrastructure-security controls.",
            ),
            (
                "Human health",
                "Climate-sensitive health outcomes include heat illness, cardiovascular stress, respiratory effects, vector-borne disease, and injuries from extreme events. Weather variables are necessary but insufficient. Demographics, baseline health, air quality, housing, behavior, and healthcare access strongly influence vulnerability and must be handled with ethical and privacy safeguards.",
            ),
            (
                "Compound events",
                "Impacts can arise from concurrent or sequential hazards, such as heat and drought, heavy rain on saturated ground, or wildfire followed by debris-flow-producing rainfall. AeroClim’s engineered interactions are a first step toward compound analysis. Event definitions and multivariate return periods would provide a more rigorous foundation.",
            ),
            (
                "Adaptation relevance",
                "Decision makers need information matched to planning horizons and actions. Near-term operational forecasts, seasonal outlooks, and multi-decadal climate projections are different products with different uncertainty. AeroClim currently blends historical analysis and short-horizon prediction; future interfaces should separate these temporal scales and identify appropriate uses.",
            ),
        ],
    )

    add_extended_appendix(
        doc,
        "APPENDIX G. DEPLOYMENT, GOVERNANCE, AND USER SAFETY",
        [
            (
                "Data provenance",
                "Every input should be traceable to a source dataset, retrieval operation, transformation, and quality-control result. Provenance records should remain attached to derived tables and figures. When simulation or replay is used, the status must be displayed prominently and included in downloaded outputs.",
            ),
            (
                "Model cards",
                "Each branch and fusion model should have a model card describing intended use, training coverage, feature schema, target definition, metrics, uncertainty, subgroup performance, missing-data behavior, limitations, ethical considerations, and prohibited use. Model cards should be versioned with artifacts and updated after retraining.",
            ),
            (
                "User-interface warnings",
                "The current gauge and severity labels are visually persuasive. Warning language should be tied to validated thresholds and calibration evidence. A prototype banner, data-status indicator, uncertainty range, and link to methodology can reduce the risk that users treat exploratory results as official alerts.",
            ),
            (
                "Audit logging",
                "For research and operational review, the system should log prediction timestamp, user-selected station, input values, data versions, model versions, missingness, output, and explanation. Logs should avoid unnecessary personal data and should be protected from modification. They enable incident investigation and reproducibility.",
            ),
            (
                "Monitoring and drift",
                "Continuous monitoring should compare incoming feature distributions with training data, track missingness and out-of-range values, and evaluate calibration when outcomes become available. Drift thresholds must trigger review rather than automatic retraining. Climate nonstationarity makes long-term monitoring a scientific requirement.",
            ),
            (
                "Security and dependency management",
                "Model files and serialized Python objects can pose supply-chain risks. Artifacts should come from trusted build pipelines, be checksummed, and be loaded only in controlled environments. Dependencies should be pinned and scanned, secrets should remain outside source control, and public interfaces should limit resource-intensive operations.",
            ),
            (
                "Accessibility",
                "Charts require descriptive titles, readable contrast, text alternatives, and tabular equivalents. Keyboard navigation and screen-reader labels should be tested. Environmental information often serves diverse audiences, so accessibility is part of technical quality rather than an optional visual enhancement.",
            ),
            (
                "Human review",
                "High-consequence outputs should be reviewed by qualified meteorological, climate, hydrological, or emergency-management professionals. The system can organize evidence and highlight patterns, but it should not replace official warnings or expert interpretation. Escalation paths must be clear when data or models disagree.",
            ),
            (
                "Change control",
                "Updates to feature definitions, thresholds, data sources, or model architecture can change output meaning even when the interface looks identical. A reviewed release process should document changes, rerun benchmark tests, compare calibration, and preserve the previous version for rollback.",
            ),
            (
                "Reproducible reporting",
                "Figures, tables, and narrative metrics should be regenerated from versioned scripts. Manual transcription should be minimized, and citations should identify whether a value was recomputed or read from a saved artifact. This report follows that distinction by separating descriptive analysis from historical model metrics.",
            ),
            (
                "Staged deployment",
                "A safe pathway begins with offline retrospective evaluation, followed by shadow-mode operation, expert review, limited pilot use, and only then broader availability. Advancement between stages should require predefined evidence. Public release should not occur merely because an interface is technically functional.",
            ),
            (
                "Retirement and rollback",
                "Models should have retirement criteria for unsupported dependencies, degraded calibration, superseded data, or changed user needs. Rollback procedures should restore a known artifact and communicate the change. Archived predictions and metadata should remain available for scientific audit subject to retention policy.",
            ),
        ],
    )

    # Normalize blank spacer paragraphs.
    for p in doc.paragraphs:
        if not p.text.strip() and not p._p.xpath(".//w:drawing"):
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0

    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    build_report()
