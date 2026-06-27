from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "report_output" / "NTCC_Term_Report_Updated.docx"

TITLE = "A Multimodal Approach to Assess Climate Change Impacts on the U.S. Environment"
AUTHORS = "Kushagra Singh (A2305224165), Chahat (A2305224162)"
AFFILIATION = "B.Tech CSE, Amity University, Noida | Academic Year: 2024-2028"
EMAIL = "Corresponding author: kushagra.singh14@s.amity.edu"


def load_json(rel: str) -> dict:
    path = ROOT / rel
    if not path.exists():
        return {}
    return json.loads(path.read_text(encoding="utf-8"))


def pct(value: float | None, digits: int = 1) -> str:
    if value is None:
        return "N/A"
    return f"{value * 100:.{digits}f}%"


def num(value: float | None, digits: int = 3) -> str:
    if value is None:
        return "N/A"
    return f"{value:.{digits}f}"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[min(idx, len(widths_dxa) - 1)])


def set_cell_margins(table, top=80, start=120, bottom=80, end=120) -> None:
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for m_name, m_value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{m_name}"))
        if node is None:
            node = OxmlElement(f"w:{m_name}")
            margins.append(node)
        node.set(qn("w:w"), str(m_value))
        node.set(qn("w:type"), "dxa")


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths_dxa: list[int], caption: str | None = None):
    if caption:
        p = doc.add_paragraph(caption)
        p.style = "CaptionText"
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_geometry(table, widths_dxa)
    set_cell_margins(table)
    hdr = table.rows[0].cells
    for idx, text in enumerate(headers):
        hdr[idx].text = text
        hdr[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(hdr[idx], "F2F4F7")
        for p in hdr[idx].paragraphs:
            p.paragraph_format.space_after = Pt(0)
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].text = text
            cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.TOP
            for p in cells[idx].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.1
                for r in p.runs:
                    r.font.size = Pt(9)
    doc.add_paragraph()
    return table


def apply_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.line_spacing = 1.1

    if "CaptionText" not in styles:
        styles.add_style("CaptionText", 1)
    cap = styles["CaptionText"]
    cap.font.name = "Calibri"
    cap.font.size = Pt(9)
    cap.font.italic = True
    cap.font.color.rgb = RGBColor(80, 80, 80)
    cap.paragraph_format.space_before = Pt(4)
    cap.paragraph_format.space_after = Pt(4)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.text = "NTCC Term Report | Multimodal Weather Forecasting"
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = RGBColor(80, 80, 80)


def add_title_block(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(TITLE)
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor.from_string("0B2545")
    p.paragraph_format.space_after = Pt(6)

    for text in [AUTHORS, AFFILIATION, EMAIL]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(60, 60, 60)
        p.paragraph_format.space_after = Pt(3)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(6)
        p.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(6)
        p.add_run(item)


def add_manual_numbered(doc: Document, items: list[str]) -> None:
    for idx, item in enumerate(items, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(6)
        p.add_run(f"{idx}.").bold = False
        p.add_run(f"\t{item}")


def build_report() -> Path:
    metrics = load_json("saved_models/metrics.json")
    cnn = load_json("saved_models/cnn_event_metrics.json")
    fusion = load_json("saved_models/cnn_tabular_fusion_metrics.json")
    relabel = load_json("data/mrms_historical_images/relabel_summary.json")
    mrms = load_json("data/mrms_historical_images/summary.json")

    doc = Document()
    apply_styles(doc)
    doc.core_properties.title = TITLE
    doc.core_properties.author = "Kushagra Singh; Chahat"
    doc.core_properties.subject = "NTCC term report updated with current trained model values"

    add_title_block(doc)

    doc.add_heading("1. Abstract", level=1)
    doc.add_paragraph(
        "This project develops a multimodal weather and climate-impact forecasting system for selected U.S. weather stations. "
        "The current pipeline combines station observations, NOAA-derived atmospheric variables, ERA5-Land soil moisture, sea-surface temperature indicators, "
        "and MRMS two-dimensional radar image chips. The goal is not only to predict future temperature patterns, but also to estimate weather-event risk by combining "
        "time-series learning, tabular machine learning, radar image classification, and late fusion. In the latest trained artifacts, the LSTM temperature model achieved "
        f"an R² of {num(metrics.get('lstm_r2'))}, MAE of {num(metrics.get('lstm_mae_c'), 2)} °C, and RMSE of {num(metrics.get('lstm_rmse_c'), 2)} °C. "
        f"The standalone MRMS CNN reached {pct(cnn.get('test_accuracy'))} accuracy and {num(cnn.get('test_auc'))} AUC, while the CNN + tabular fusion model improved event detection to "
        f"{pct(fusion.get('test_accuracy'))} accuracy, {num(fusion.get('test_f1'))} F1, and {num(fusion.get('test_auc'))} AUC. These values show that the strongest current event signal comes from "
        "combining spatial radar evidence with station/tabular features."
    )

    doc.add_heading("2. Introduction", level=1)
    doc.add_paragraph(
        "Climate change assessment increasingly requires models that can combine long-term climate trends with short-term weather-event indicators. Traditional station datasets are valuable for "
        "temperature, precipitation, wind, pressure, humidity, and derived climate variables, but they do not directly capture the spatial storm structure visible in radar. This report therefore "
        "updates the original term-paper draft into a project-specific technical report for a multimodal system that uses both tabular time-series data and MRMS 2D radar images."
    )
    doc.add_paragraph(
        "The current system focuses on a reduced 20-station dataset so the pipeline can remain practical on a local machine while still covering geographically different U.S. environments. "
        "The architecture includes LSTM forecasting for continuous station weather, Random Forest models for tabular event-risk estimation, a 2D CNN for MRMS radar-image event classification, "
        "and fusion artifacts that combine complementary model outputs."
    )

    doc.add_heading("3. Objectives", level=1)
    add_bullets(doc, [
        "Build a clean, reproducible dataset using selected U.S. weather stations and aligned atmospheric, soil-moisture, SST, and radar-image features.",
        "Train and compare LSTM, Random Forest, CNN, and fusion models for weather forecasting and event-risk prediction.",
        "Use real MRMS 2D image chips from RadarOnly_QPE_01H, MergedReflectivityQCComposite, and MergedAzShear_0-2kmAGL as image channels for CNN training.",
        "Replace earlier placeholder report values with the latest saved model metrics from the project artifacts.",
        "Identify current limitations and future improvements, especially label alignment, event balance, and longer historical radar coverage.",
    ])

    doc.add_heading("4. Literature Review", level=1)
    literature_rows = [
        ["1", "Hochreiter and Schmidhuber", "1997", "Sequence modeling", "Introduced LSTM networks, which are suitable for weather time series because they preserve temporal memory."],
        ["2", "Breiman", "2001", "Random forests", "Presented ensemble decision trees that remain useful for tabular climate and station-feature modeling."],
        ["3", "LeCun et al.", "1998", "Image recognition", "Established CNN concepts used here for learning spatial structure from radar image chips."],
        ["4", "NOAA GSOD/NCEI", "Ongoing", "Station weather records", "Provides daily station-scale observations used for temperature and tabular forecasting."],
        ["5", "NOAA Storm Events Database", "Ongoing", "Storm event reports", "Supplies reported event labels used to align MRMS images with weather-event outcomes."],
        ["6", "NOAA MRMS", "Ongoing", "2D radar products", "Provides gridded radar/QPE/azimuthal-shear fields used as CNN image channels."],
        ["7", "ECMWF ERA5-Land", "Ongoing", "Soil moisture reanalysis", "Adds land-surface moisture context that can influence heat, precipitation, and storm potential."],
        ["8", "Climate index research", "Multiple", "ENSO/NAO/PDO indicators", "Shows that large-scale indices can improve long-range climate interpretation."],
        ["9", "ConvLSTM studies", "2015+", "Spatiotemporal forecasting", "Demonstrates that spatial and temporal weather signals can be modeled jointly when larger image sequences are available."],
        ["10", "Remote-sensing fusion studies", "2018+", "Multimodal Earth observation", "Supports combining imagery with tabular environmental variables for stronger predictions."],
        ["11", "Extreme-weather ML studies", "2020+", "Event classification", "Highlights the importance of class balance, chronological validation, and event-window design."],
        ["12", "Hydrometeorological modeling", "2020+", "Precipitation and soil data", "Shows that soil moisture and precipitation history can improve practical risk estimation."],
        ["13", "Deep learning forecast systems", "2021+", "WeatherBench-style datasets", "Provides evidence that neural models need large, well-aligned weather archives to generalize."],
        ["14", "Explainable AI for climate", "2021+", "Feature importance/Grad-CAM", "Motivates interpretable RF importance and radar saliency checks for project reporting."],
        ["15", "Multimodal fusion literature", "2022+", "Late fusion", "Supports blending independent model probabilities when modalities have different noise patterns."],
        ["16", "Operational nowcasting studies", "2022+", "Radar nowcasting", "Shows why MRMS radar is important for short-term event detection."],
        ["17", "Climate-risk dashboards", "2023+", "Decision support", "Connects model outputs to practical dashboards for planning and risk communication."],
    ]
    add_table(
        doc,
        ["Ref.", "Author / Source", "Year", "Dataset / Method", "Small Analysis"],
        literature_rows[:9],
        [700, 1900, 850, 1900, 4010],
        "Table 4.1(a): Summary of literature and data sources relevant to the project.",
    )
    add_table(
        doc,
        ["Ref.", "Author / Source", "Year", "Dataset / Method", "Small Analysis"],
        literature_rows[9:],
        [700, 1900, 850, 1900, 4010],
        "Table 4.1(b): Continued literature and data-source summary.",
    )

    doc.add_heading("5. Dataset and Methodology", level=1)
    doc.add_paragraph(
        "The dataset is organized around daily station records and aligned auxiliary environmental variables. The station dataset has been reduced from the earlier 96-station setup to 20 selected stations, "
        "which reduces training cost while keeping enough station diversity for experimentation. Atmospheric features are fetched from NOAA sources, soil moisture is fetched from Copernicus/ERA5-Land, "
        "and radar imagery is prepared from NOAA MRMS 2D products."
    )
    add_bullets(doc, [
        "Tabular modalities: station weather observations, atmospheric variables, ERA5-Land soil moisture, SST indicators, and engineered calendar/lag features.",
        "Image modalities: 64 × 64 × 3 MRMS image chips where the three channels are MergedReflectivityQCComposite, RadarOnly_QPE_01H, and MergedAzShear_0-2kmAGL.",
        "Temporal validation: event-image models use a chronological 2022/2023/2024 split to better reflect future-year generalization.",
        "Event labels: NOAA Storm Events are aligned to MRMS samples using a ±6 hour event window, with a 12 hour ambiguity zone excluded from training.",
    ])
    add_table(
        doc,
        ["Dataset component", "Current value", "Role in model"],
        [
            ["Selected station scope", "20 U.S. stations", "Keeps the local pipeline faster while preserving regional variety."],
            ["MRMS total chips", str(relabel.get("samples", mrms.get("samples", "N/A"))), "Raw radar-image examples prepared for CNN/fusion training."],
            ["MRMS labeled chips", str(relabel.get("labeled_samples", "N/A")), "Samples retained after event-window labeling and ambiguity filtering."],
            ["Quiet/event labels", f"{relabel.get('label_counts', {}).get('0', 'N/A')} quiet / {relabel.get('label_counts', {}).get('1', 'N/A')} event", "Shows the class imbalance that affects CNN learning."],
            ["Ambiguous labels removed", str(relabel.get("label_counts", {}).get("-1", "N/A")), "Excluded to avoid uncertain supervision near event boundaries."],
            ["Image shape", "64 × 64 × 3", "CNN input tensor using three MRMS radar products."],
        ],
        [2200, 2300, 4860],
        "Table 5.1: Current dataset status after station reduction and MRMS relabeling.",
    )

    doc.add_heading("6. Analysis and Key Findings", level=1)
    doc.add_heading("6.1 Individual Model Performance Analysis", level=2)
    add_table(
        doc,
        ["Model", "Input type", "Main task", "Current result", "Interpretation"],
        [
            ["LSTM", "Station time-series features", "Next-step temperature/weather forecasting", f"R² {num(metrics.get('lstm_r2'))}; MAE {num(metrics.get('lstm_mae_c'), 2)} °C; RMSE {num(metrics.get('lstm_rmse_c'), 2)} °C", "Strongest for continuous temperature patterns because it learns sequence dependence."],
            ["Random Forest", "Tabular climate/atmospheric features", "Broad hazard classification", f"Accuracy {pct(metrics.get('rf_accuracy'))}; AUC {num(metrics.get('rf_auc'))}", "Useful baseline, but the broad proxy target is noisy and only modestly separable."],
            ["Event Random Forest", "Station/tabular features with event labels", "NOAA event-risk classification", f"Accuracy {pct(metrics.get('event_model_test_accuracy'))}; F1 {num(metrics.get('event_model_test_f1'))}; AUC {num(metrics.get('event_model_test_auc'))}", "Much stronger AUC when trained against explicit storm-event labels."],
            ["2D CNN", "MRMS radar image chips", "Radar-based event classification", f"Accuracy {pct(cnn.get('test_accuracy'))}; F1 {num(cnn.get('test_f1'))}; AUC {num(cnn.get('test_auc'))}", "Learns spatial radar signatures, but performance is limited by sample count and class imbalance."],
        ],
        [1450, 1850, 1800, 2050, 2210],
        "Table 6.1: Updated individual model performance from saved artifacts.",
    )

    doc.add_heading("6.2 LSTM", level=2)
    doc.add_paragraph(
        "The LSTM is used for sequential station forecasting. It learns how recent weather conditions influence the next forecast step. In the current artifacts, its R² of "
        f"{num(metrics.get('lstm_r2'))} and MAE of {num(metrics.get('lstm_mae_c'), 2)} °C indicate that it is the best-performing model for continuous temperature-style prediction. "
        "However, LSTM alone is not an image model and does not directly read radar structure."
    )

    doc.add_heading("6.3 Random Forest", level=2)
    doc.add_paragraph(
        "The Random Forest models use tabular variables and are valuable because they are fast, robust, and easier to interpret than deep networks. The general RF hazard proxy produced "
        f"{pct(metrics.get('rf_accuracy'))} accuracy and {num(metrics.get('rf_auc'))} AUC, while the event-specific RF achieved {num(metrics.get('event_model_test_auc'))} AUC. "
        "This contrast shows that label quality is more important than model complexity: a clearer event target gives much better separation."
    )

    doc.add_heading("6.4 2D CNN", level=2)
    doc.add_paragraph(
        "The 2D CNN uses MRMS radar images as three-channel inputs. It is appropriate here because radar products are spatial grids, similar to images, where storm shape, intensity, precipitation, "
        "and rotation signatures matter. The current CNN used 2,129 labeled samples with a chronological 2022/2023/2024 split and achieved "
        f"{pct(cnn.get('test_accuracy'))} accuracy, {num(cnn.get('test_f1'))} F1, and {num(cnn.get('test_auc'))} AUC. This is a working proof of concept, but it needs more balanced event and quiet dates."
    )

    doc.add_heading("6.5 Multimodal Fusion", level=2)
    doc.add_paragraph(
        "Fusion combines predictions from different model families. This is useful because the models observe different evidence: the CNN sees spatial radar patterns, while tabular models see station, "
        "atmospheric, land-surface, and climate context. The strongest current event artifact is the CNN + tabular fusion model."
    )
    add_table(
        doc,
        ["Fusion artifact", "Inputs combined", "Current result", "Meaning"],
        [
            ["LSTM/RF fusion", "Sequence forecast + tabular model", f"Accuracy {pct(metrics.get('fusion_accuracy'))}; AUC {num(metrics.get('fusion_auc'))}", "Works as a baseline but does not yet deliver strong event separation."],
            ["CNN + tabular fusion", "CNN event probability + tabular event probability", f"Accuracy {pct(fusion.get('test_accuracy'))}; F1 {num(fusion.get('test_f1'))}; AUC {num(fusion.get('test_auc'))}; AP {num(fusion.get('test_average_precision'))}", "Best current event model because it blends radar image evidence with contextual tabular risk."],
        ],
        [1800, 2600, 2500, 2460],
        "Table 6.2: Updated fusion-model performance.",
    )

    doc.add_heading("6.6 Ablation Study", level=2)
    add_table(
        doc,
        ["Configuration", "What it removes", "Observed effect"],
        [
            ["LSTM only", "No radar or event labels", "Good continuous forecasting, but not designed for direct storm-event classification."],
            ["Random Forest broad proxy", "No spatial radar images", "Fast and interpretable, but broad hazard labels produce weak AUC."],
            ["CNN only", "No station/tabular context", "Uses spatial storm structure, but limited by 2,129 labeled samples and imbalance."],
            ["CNN + tabular fusion", "No modality removed", "Best current event result: AUC 0.853 and F1 0.582 on the 2024 test split."],
        ],
        [2100, 2700, 4560],
        "Table 6.3: Practical ablation interpretation from the current trained artifacts.",
    )

    doc.add_heading("6.7 Key Findings", level=2)
    add_bullets(doc, [
        f"LSTM is currently the strongest continuous weather forecaster, with R² {num(metrics.get('lstm_r2'))} and MAE {num(metrics.get('lstm_mae_c'), 2)} °C.",
        f"The standalone CNN is implemented and working, but its AUC of {num(cnn.get('test_auc'))} shows that more balanced MRMS dates are needed.",
        f"The CNN + tabular fusion model is the strongest event-risk artifact, reaching AUC {num(fusion.get('test_auc'))} and F1 {num(fusion.get('test_f1'))}.",
        "The reduced 20-station setup is more practical for iteration, but final-grade generalization would benefit from a larger station/date archive.",
        "Future improvement should prioritize better event-window alignment, more quiet/event-balanced MRMS sampling, and longer historical radar coverage.",
    ])

    doc.add_heading("7. Applications", level=1)
    add_numbered(doc, [
        "Climate trend forecasting: estimate changing temperature and precipitation patterns for selected U.S. regions.",
        "Extreme weather detection: classify event risk using station features, NOAA events, and radar image structure.",
        "Climate-risk assessment and planning: support city or regional planning with interpretable model outputs.",
        "Disaster preparedness: provide early warning support by combining tabular and radar signals.",
        "Agriculture and water management: use temperature, precipitation, and soil-moisture indicators for operational decisions.",
        "Policy dashboards: communicate station-level risk, model confidence, and climate drivers to non-technical users.",
        "Scientific research: compare LSTM, RF, CNN, and fusion behavior across climate/weather modalities.",
    ])

    doc.add_heading("8. Limitations and Future Work", level=1)
    add_bullets(doc, [
        "Increase MRMS training coverage from the current small historical sample toward hundreds or thousands of balanced dates.",
        "Improve label alignment by matching storm reports to radar grid location, event type, and observation time more precisely.",
        "Tune the CNN architecture with stronger augmentation, focal loss, and calibration once a larger dataset is available.",
        "Add ConvLSTM or CNN-LSTM models if sequential radar-image windows are downloaded, because static 2D CNNs only see one timestamp at a time.",
        "Run full cross-station and cross-year validation before presenting the system as operationally reliable.",
    ])

    doc.add_heading("9. Conclusion", level=1)
    doc.add_paragraph(
        "The updated project now implements all three major model families discussed during development: LSTM for sequential station forecasting, Random Forest for tabular risk estimation, and 2D CNN "
        "for MRMS radar-image classification. Fusion artifacts are also implemented, including a CNN + tabular fusion model that currently gives the best event-risk performance. The most important "
        "next step is data quality rather than simply adding model layers: more balanced MRMS dates, sharper event labels, and larger historical coverage should improve the CNN and fusion results."
    )

    doc.add_heading("10. References", level=1)
    references = [
        "Hochreiter, S., and Schmidhuber, J. (1997). Long Short-Term Memory.",
        "Breiman, L. (2001). Random Forests.",
        "LeCun, Y., Bottou, L., Bengio, Y., and Haffner, P. (1998). Gradient-based learning applied to document recognition.",
        "NOAA National Centers for Environmental Information. Global Summary of the Day and Storm Events Database.",
        "NOAA National Severe Storms Laboratory. Multi-Radar/Multi-Sensor System (MRMS) 2D products.",
        "ECMWF Copernicus Climate Data Store. ERA5-Land reanalysis and soil-moisture variables.",
        "Recent literature on radar nowcasting, multimodal remote-sensing fusion, and machine-learning-based climate-risk modeling.",
    ]
    add_manual_numbered(doc, references)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build_report())
